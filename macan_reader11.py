import sys
import os
import pypdfium2 as pdfium
import io

from PySide6.QtWidgets import (
    QApplication, QMainWindow, QWidget, QVBoxLayout, QHBoxLayout, QLabel,
    QScrollArea, QSplitter, QFileDialog, QStatusBar, QSlider, QMessageBox,
    QToolBar, QPushButton, QLineEdit, QCheckBox, QMenu,
    QInputDialog,
    QTabWidget,  # <-- BARU
    QTabBar      # <-- BARU
)
from PySide6.QtGui import (
    QPixmap, QImage, QIcon, QAction, QPainter, QScreen, QColor,
    QActionGroup,
    QFont, QFontMetrics
)
from PySide6.QtCore import (
    Qt, QSize, QSettings, QUrl, QByteArray, Signal, QRectF,
    QPoint # <-- BARU
)
from PySide6.QtSvg import QSvgRenderer
from PySide6.QtPrintSupport import QPrinter, QPrintPreviewDialog

# --- Helper untuk menangani path saat di-bundle dengan PyInstaller ---
def resource_path(relative_path):
    """ Dapatkan path absolut ke resource, berfungsi untuk dev dan PyInstaller """
    try:
        base_path = sys._MEIPASS
    except Exception:
        base_path = os.path.abspath(".")
    return os.path.join(base_path, relative_path)

# --- Icon SVG dalam format XML (Embedded) ---
SVG_ICONS = {
    "open": """<svg xmlns="http://www.w3.org/2000/svg" width="24" height="24" viewBox="0 0 24 24" fill="none" stroke="currentColor" stroke-width="2" stroke-linecap="round" stroke-linejoin="round"><path d="M22 19a2 2 0 0 1-2 2H4a2 2 0 0 1-2-2V5a2 2 0 0 1 2-2h5l2 3h9a2 2 0 0 1 2 2z"></path></svg>""",
    "print": """<svg xmlns="http://www.w3.org/2000/svg" width="24" height="24" viewBox="0 0 24 24" fill="none" stroke="currentColor" stroke-width="2" stroke-linecap="round" stroke-linejoin="round"><polyline points="6 9 6 2 18 2 18 9"></polyline><path d="M6 18H4a2 2 0 0 1-2-2v-5a2 2 0 0 1 2-2h16a2 2 0 0 1 2 2v5a2 2 0 0 1-2 2h-2"></path><rect x="6" y="14" width="12" height="8"></rect></svg>""",
    "zoom-in": """<svg xmlns="http://www.w3.org/2000/svg" width="24" height="24" viewBox="0 0 24 24" fill="none" stroke="currentColor" stroke-width="2" stroke-linecap="round" stroke-linejoin="round"><circle cx="11" cy="11" r="8"></circle><line x1="21" y1="21" x2="16.65" y2="16.65"></line><line x1="11" y1="8" x2="11" y2="14"></line><line x1="8" y1="11" x2="14" y2="11"></line></svg>""",
    "zoom-out": """<svg xmlns="http://www.w3.org/2000/svg" width="24" height="24" viewBox="0 0 24 24" fill="none" stroke="currentColor" stroke-width="2" stroke-linecap="round" stroke-linejoin="round"><circle cx="11" cy="11" r="8"></circle><line x1="21" y1="21" x2="16.65" y2="16.65"></line><line x1="8" y1="11" x2="14" y2="11"></line></svg>""",
    "zoom-reset": """<svg xmlns="http://www.w3.org/2000/svg" width="24" height="24" viewBox="0 0 24 24" fill="none" stroke="currentColor" stroke-width="2" stroke-linecap="round" stroke-linejoin="round"><circle cx="12" cy="12" r="10"></circle><circle cx="12" cy="12" r="6"></circle><circle cx="12" cy="12" r="2"></circle></svg>""",
    "prev-page": """<svg xmlns="http://www.w3.org/2000/svg" width="24" height="24" viewBox="0 0 24 24" fill="none" stroke="currentColor" stroke-width="2" stroke-linecap="round" stroke-linejoin="round"><polyline points="15 18 9 12 15 6"></polyline></svg>""",
    "next-page": """<svg xmlns="http://www.w3.org/2000/svg" width="24" height="24" viewBox="0 0 24 24" fill="none" stroke="currentColor" stroke-width="2" stroke-linecap="round" stroke-linejoin="round"><polyline points="9 18 15 12 9 6"></polyline></svg>""",
    "jump": """<svg xmlns="http://www.w3.org/2000/svg" width="24" height="24" viewBox="0 0 24 24" fill="none" stroke="currentColor" stroke-width="2" stroke-linecap="round" stroke-linejoin="round"><line x1="4" y1="9" x2="20" y2="9"></line><line x1="4" y1="15" x2="20" y2="15"></line><line x1="10" y1="3" x2="8" y2="21"></line><line x1="16" y1="3" x2="14" y2="21"></line></svg>""",
    "save": """<svg xmlns="http://www.w3.org/2000/svg" width="24" height="24" viewBox="0 0 24 24" fill="none" stroke="currentColor" stroke-width="2" stroke-linecap="round" stroke-linejoin="round"><path d="M19 21H5a2 2 0 0 1-2-2V5a2 2 0 0 1 2-2h11l5 5v11a2 2 0 0 1-2 2z"></path><polyline points="17 21 17 13 7 13 7 21"></polyline><polyline points="7 3 7 8 15 8"></polyline></svg>""",
    "search": """<svg xmlns="http://www.w3.org/2000/svg" width="24" height="24" viewBox="0 0 24 24" fill="none" stroke="currentColor" stroke-width="2" stroke-linecap="round" stroke-linejoin="round"><circle cx="11" cy="11" r="8"></circle><line x1="21" y1="21" x2="16.65" y2="16.65"></line></svg>""",
    "hand": """<svg xmlns="http://www.w3.org/2000/svg" width="24" height="24" viewBox="0 0 24 24" fill="none" stroke="currentColor" stroke-width="2" stroke-linecap="round" stroke-linejoin="round"><path d="M18 11V6a2 2 0 0 0-2-2v0a2 2 0 0 0-2 2v0"></path><path d="M14 10V4a2 2 0 0 0-2-2v0a2 2 0 0 0-2 2v2"></path><path d="M10 10.5V6a2 2 0 0 0-2-2v0a2 2 0 0 0-2 2v8"></path><path d="M18 8a2 2 0 0 1 2 2v10a2 2 0 0 1-2 2h-4a2 2 0 0 1-2-2v-4a2 2 0 0 1 2-2h2"></path></svg>""",
    # --- ICON BARU UNTUK FULLSCREEN ---
    "fullscreen-enter": """<svg xmlns="http://www.w3.org/2000/svg" width="24" height="24" viewBox="0 0 24 24" fill="none" stroke="currentColor" stroke-width="2" stroke-linecap="round" stroke-linejoin="round"><path d="M8 3H5a2 2 0 0 0-2 2v3m18 0V5a2 2 0 0 0-2-2h-3m0 18h3a2 2 0 0 0 2-2v-3M3 16v3a2 2 0 0 0 2 2h3"></path></svg>""",
    "fullscreen-exit": """<svg xmlns="http://www.w3.org/2000/svg" width="24" height="24" viewBox="0 0 24 24" fill="none" stroke="currentColor" stroke-width="2" stroke-linecap="round" stroke-linejoin="round"><path d="M8 3v3a2 2 0 0 1-2 2H3m18 0h-3a2 2 0 0 1-2-2V3m0 18v-3a2 2 0 0 0-2-2h-3M3 16h3a2 2 0 0 0 2 2v3"></path></svg>"""
}

# Kamus untuk warna ikon tema
THEME_ICON_COLORS = {
    "light": "black",
    "dark": "#E0E0E0",
    "dark-blue": "#E0E0E0",
    "neon-blue": "#E0E0E0",
    "soft-pink": "#333333"
}

# Kamus untuk QSS Tema
THEME_STYLESHEETS = {
    "light": """
        QWidget {
            background-color: #FFFFFF;
            color: #000000;
        }
        QMainWindow, QToolBar, QStatusBar, QMenu, QMenuBar {
            background-color: #F0F0F0;
            color: #000000;
        }
        QMenuBar::item:selected, QMenu::item:selected {
            background-color: #0078D4;
            color: #FFFFFF;
        }
        QScrollArea {
            background-color: #E0E0E0;
            border: 1px solid #C0C0C0;
        }
        QPushButton {
            background-color: #E1E1E1;
            border: 1px solid #C0C0C0;
            padding: 5px;
        }
        QPushButton:hover {
            background-color: #E5E5E5;
        }
        QPushButton:pressed {
            background-color: #D0D0D0;
        }
        /* Style untuk tombol status bar */
        QStatusBar QPushButton {
            padding: 2px;
            border-radius: 4px;
        }
        QStatusBar QPushButton:hover { background-color: #E5E5E5; }
        QStatusBar QPushButton:pressed { background-color: #D0D0D0; }
        QStatusBar QPushButton:checked { background-color: #C0E0FF; }
        
        QLineEdit {
            background-color: #FFFFFF;
            border: 1px solid #C0C0C0;
        }
        QSlider::groove:horizontal {
            background: #C0C0C0;
            height: 8px;
            border-radius: 4px;
        }
        QSlider::handle:horizontal {
            background: #0078D4;
            width: 18px;
            height: 18px;
            border-radius: 9px;
            margin: -5px 0;
        }
        QSplitter::handle {
            background-color: #C0C0C0;
        }
        QWidget#searchWidget, QWidget#thumbnailTitleBar {
            background-color: #F0F0F0;
        }
        QWidget#thumbnailContainer {
            background-color: #E0E0E0;
        }
        ThumbnailLabel {
            border: 1px solid lightgray;
            padding: 2px;
            margin: 5px;
        }
        ThumbnailLabel:hover {
            border: 1px solid dodgerblue;
        }
        /* Style untuk Tab */
        QTabWidget::pane { border-top: 1px solid #C0C0C0; }
        QTabBar::tab {
            background: #E1E1E1; color: #555; padding: 8px 15px;
            border: 1px solid #C0C0C0; border-bottom: none;
            border-top-left-radius: 5px; border-top-right-radius: 5px;
            margin-left: 2px;
        }
        QTabBar::tab:selected { background: #FFFFFF; color: #111; }
        QTabBar::tab:!selected:hover { background: #DCDCDC; }
    """,
    "dark": """
        QWidget {
            background-color: #3C3C3C;
            color: #E0E0E0;
        }
        QMainWindow, QToolBar, QStatusBar, QMenu, QMenuBar {
            background-color: #2D2D2D;
            color: #E0E0E0;
        }
        QMenuBar::item:selected, QMenu::item:selected {
            background-color: #0078D4;
            color: #FFFFFF;
        }
        QScrollArea {
            background-color: #333333;
            border: 1px solid #4A4A4A;
        }
        QPushButton {
            background-color: #4A4A4A;
            border: 1px solid #5A5A5A;
            padding: 5px;
        }
        QPushButton:hover {
            background-color: #5A5A5A;
        }
        QPushButton:pressed {
            background-color: #444444;
        }
        /* Style untuk tombol status bar */
        QStatusBar QPushButton {
            padding: 2px;
            border-radius: 4px;
        }
        QStatusBar QPushButton:hover { background-color: #5A5A5A; }
        QStatusBar QPushButton:pressed { background-color: #444444; }
        QStatusBar QPushButton:checked { background-color: #005A9E; }
        
        QLineEdit {
            background-color: #333333;
            border: 1px solid #5A5A5A;
        }
        QSlider::groove:horizontal {
            background: #4A4A4A;
            height: 8px;
            border-radius: 4px;
        }
        QSlider::handle:horizontal {
            background: #0078D4;
            width: 18px;
            height: 18px;
            border-radius: 9px;
            margin: -5px 0;
        }
        QSplitter::handle {
            background-color: #4A4A4A;
        }
        QWidget#searchWidget, QWidget#thumbnailTitleBar {
            background-color: #2D2D2D;
        }
        QWidget#thumbnailContainer {
            background-color: #333333;
        }
        ThumbnailLabel {
            border: 1px solid #5A5A5A;
            padding: 2px;
            margin: 5px;
        }
        ThumbnailLabel:hover {
            border: 1px solid #0078D4;
        }
        /* Style untuk Tab */
        QTabWidget::pane { border-top: 1px solid #4A4A4A; }
        QTabBar::tab {
            background: #3C3C3C; color: #AAA; padding: 8px 15px;
            border: 1px solid #4A4A4A; border-bottom: none;
            border-top-left-radius: 5px; border-top-right-radius: 5px;
            margin-left: 2px;
        }
        QTabBar::tab:selected { background: #333333; color: #EEE; }
        QTabBar::tab:!selected:hover { background: #4A4A4A; }
    """,
    "dark-blue": """
        QWidget {
            background-color: #2D3A4F;
            color: #E0E0E0;
        }
        QMainWindow, QToolBar, QStatusBar, QMenu, QMenuBar {
            background-color: #222B3A;
            color: #E0E0E0;
        }
        QMenuBar::item:selected, QMenu::item:selected {
            background-color: #4A90E2;
            color: #FFFFFF;
        }
        QScrollArea {
            background-color: #283344;
            border: 1px solid #3A4A63;
        }
        QPushButton {
            background-color: #3A4A63;
            border: 1px solid #4A5B79;
            padding: 5px;
        }
        QPushButton:hover {
            background-color: #4A5B79;
        }
        QPushButton:pressed {
            background-color: #33425A;
        }
        /* Style untuk tombol status bar */
        QStatusBar QPushButton {
            padding: 2px;
            border-radius: 4px;
        }
        QStatusBar QPushButton:hover { background-color: #4A5B79; }
        QStatusBar QPushButton:pressed { background-color: #33425A; }
        QStatusBar QPushButton:checked { background-color: #4A90E2; }
        
        QLineEdit {
            background-color: #283344;
            border: 1px solid #4A5B79;
        }
        QSlider::groove:horizontal {
            background: #3A4A63;
            height: 8px;
            border-radius: 4px;
        }
        QSlider::handle:horizontal {
            background: #4A90E2;
            width: 18px;
            height: 18px;
            border-radius: 9px;
            margin: -5px 0;
        }
        QSplitter::handle {
            background-color: #3A4A63;
        }
        QWidget#searchWidget, QWidget#thumbnailTitleBar {
            background-color: #222B3A;
        }
        QWidget#thumbnailContainer {
            background-color: #283344;
        }
        ThumbnailLabel {
            border: 1px solid #3A4A63;
            padding: 2px;
            margin: 5px;
        }
        ThumbnailLabel:hover {
            border: 1px solid #4A90E2;
        }
        /* Style untuk Tab */
        QTabWidget::pane { border-top: 1px solid #3A4A63; }
        QTabBar::tab {
            background: #2D3A4F; color: #AAA; padding: 8px 15px;
            border: 1px solid #3A4A63; border-bottom: none;
            border-top-left-radius: 5px; border-top-right-radius: 5px;
            margin-left: 2px;
        }
        QTabBar::tab:selected { background: #283344; color: #EEE; }
        QTabBar::tab:!selected:hover { background: #3A4A63; }
    """,
    "neon-blue": """
        QWidget {
            background-color: #0A0A1A;
            color: #00FFD1;
        }
        QMainWindow, QToolBar, QStatusBar, QMenu, QMenuBar {
            background-color: #0F0F2A;
            color: #00FFD1;
            border: 1px solid #00FFD1;
        }
        QMenuBar::item:selected, QMenu::item:selected {
            background-color: #00FFD1;
            color: #000000;
        }
        QScrollArea {
            background-color: #050510;
            border: 1px solid #00FFD1;
        }
        QPushButton {
            background-color: #1A1A3A;
            border: 1px solid #00FFD1;
            color: #00FFD1;
            padding: 5px;
        }
        QPushButton:hover {
            background-color: #2A2A4A;
        }
        QPushButton:pressed {
            background-color: #11112A;
        }
        /* Style untuk tombol status bar */
        QStatusBar QPushButton {
            padding: 2px;
            border-radius: 4px;
        }
        QStatusBar QPushButton:hover { background-color: #2A2A4A; }
        QStatusBar QPushButton:pressed { background-color: #11112A; }
        QStatusBar QPushButton:checked { background-color: #00FFD1; color: #000; }
        
        QLineEdit {
            background-color: #050510;
            border: 1px solid #00FFD1;
            color: #00FFD1;
        }
        QSlider::groove:horizontal {
            background: #1A1A3A;
            height: 8px;
            border-radius: 4px;
        }
        QSlider::handle:horizontal {
            background: #00FFD1;
            width: 18px;
            height: 18px;
            border-radius: 9px;
            margin: -5px 0;
            border: 1px solid #00FFD1;
        }
        QSplitter::handle {
            background-color: #1A1A3A;
            border: 1px solid #00FFD1;
        }
        QWidget#searchWidget, QWidget#thumbnailTitleBar {
            background-color: #0F0F2A;
        }
        QWidget#thumbnailContainer {
            background-color: #050510;
        }
        ThumbnailLabel {
            border: 1px solid #00FFD1;
            padding: 2px;
            margin: 5px;
        }
        ThumbnailLabel:hover {
            border: 2px solid #00FFD1;
        }
        /* Style untuk Tab */
        QTabWidget::pane { border-top: 1px solid #00FFD1; }
        QTabBar::tab {
            background: #1A1A3A; color: #00FFD1; padding: 8px 15px;
            border: 1px solid #00FFD1; border-bottom: none;
            border-top-left-radius: 5px; border-top-right-radius: 5px;
            margin-left: 2px;
        }
        QTabBar::tab:selected { background: #0A0A1A; color: #00FFD1; }
        QTabBar::tab:!selected:hover { background: #2A2A4A; }
    """,
    "soft-pink": """
        QWidget {
            background-color: #FFF0F5;
            color: #5D4037;
        }
        QMainWindow, QToolBar, QStatusBar, QMenu, QMenuBar {
            background-color: #FFDDEE;
            color: #5D4037;
        }
        QMenuBar::item:selected, QMenu::item:selected {
            background-color: #FF80AB;
            color: #FFFFFF;
        }
        QScrollArea {
            background-color: #FFF5F9;
            border: 1px solid #FFC0CB;
        }
        QPushButton {
            background-color: #FFC0CB;
            border: 1px solid #FF80AB;
            color: #5D4037;
            padding: 5px;
        }
        QPushButton:hover {
            background-color: #FFB0C1;
        }
        QPushButton:pressed {
            background-color: #FFA0B1;
        }
        /* Style untuk tombol status bar */
        QStatusBar QPushButton {
            padding: 2px;
            border-radius: 4px;
        }
        QStatusBar QPushButton:hover { background-color: #FFB0C1; }
        QStatusBar QPushButton:pressed { background-color: #FFA0B1; }
        QStatusBar QPushButton:checked { background-color: #FF80AB; color: white; }
        
        QLineEdit {
            background-color: #FFF5F9;
            border: 1px solid #FFC0CB;
            color: #5D4037;
        }
        QSlider::groove:horizontal {
            background: #FFC0CB;
            height: 8px;
            border-radius: 4px;
        }
        QSlider::handle:horizontal {
            background: #FF80AB;
            width: 18px;
            height: 18px;
            border-radius: 9px;
            margin: -5px 0;
        }
        QSplitter::handle {
            background-color: #FFC0CB;
        }
        QWidget#searchWidget, QWidget#thumbnailTitleBar {
            background-color: #FFDDEE;
        }
        QWidget#thumbnailContainer {
            background-color: #FFF5F9;
        }
        ThumbnailLabel {
            border: 1px solid #FFC0CB;
            padding: 2px;
            margin: 5px;
        }
        ThumbnailLabel:hover {
            border: 1px solid #FF80AB;
        }
        /* Style untuk Tab */
        QTabWidget::pane { border-top: 1px solid #FFC0CB; }
        QTabBar::tab {
            background: #FFDDEE; color: #5D4037; padding: 8px 15px;
            border: 1px solid #FFC0CB; border-bottom: none;
            border-top-left-radius: 5px; border-top-right-radius: 5px;
            margin-left: 2px;
        }
        QTabBar::tab:selected { background: #FFF0F5; color: #5D4037; }
        QTabBar::tab:!selected:hover { background: #FFC0CB; }
    """
}


def create_svg_icon(svg_xml, color="black"):
    """Membuat QIcon dari string XML SVG."""
    svg_xml = svg_xml.replace('stroke="currentColor"', f'stroke="{color}"')
    renderer = QSvgRenderer(QByteArray(svg_xml.encode('utf-8')))
    pixmap = QPixmap(QSize(24, 24))
    pixmap.fill(Qt.GlobalColor.transparent)
    painter = QPainter(pixmap)
    renderer.render(painter)
    painter.end()
    return QIcon(pixmap)

class ThumbnailLabel(QLabel):
    """Label kustom untuk thumbnail yang bisa diklik."""
    page_clicked = Signal(int)
    export_requested = Signal(int, str)

    def __init__(self, page_num, parent=None):
        super().__init__(parent)
        self.page_num = page_num
        self.setToolTip(f"Go to page {self.page_num + 1}")
        self.setCursor(Qt.CursorShape.PointingHandCursor)
        self.setContextMenuPolicy(Qt.ContextMenuPolicy.CustomContextMenu)
        self.customContextMenuRequested.connect(self.show_context_menu)

    def mousePressEvent(self, event):
        if event.button() == Qt.MouseButton.LeftButton:
            self.page_clicked.emit(self.page_num)
        elif event.button() != Qt.MouseButton.RightButton:
            super().mousePressEvent(event)

    def show_context_menu(self, pos):
        context_menu = QMenu(self)
        export_menu = context_menu.addMenu("Export Page As...")
        
        export_pdf_action = export_menu.addAction("PDF...")
        export_png_action = export_menu.addAction("PNG...")
        export_docx_action = export_menu.addAction("DOCX...")
        
        action = context_menu.exec(self.mapToGlobal(pos))
        
        if action == export_pdf_action:
            self.export_requested.emit(self.page_num, "pdf")
        elif action == export_png_action:
            self.export_requested.emit(self.page_num, "png")
        elif action == export_docx_action:
            self.export_requested.emit(self.page_num, "docx")

class PdfViewLabel(QLabel):
    """Label kustom untuk tampilan PDF yang mendukung panning."""
    def __init__(self, scroll_area, parent=None):
        super().__init__(parent)
        self.scroll_area = scroll_area
        self.pan_mode = False
        self.panning = False
        self.last_pan_pos = None
        self.setCursor(Qt.CursorShape.ArrowCursor)

    def setPanMode(self, enabled):
        """Mengaktifkan atau menonaktifkan mode pan."""
        self.pan_mode = enabled
        self.setCursor(Qt.CursorShape.OpenHandCursor if enabled else Qt.CursorShape.ArrowCursor)

    def mousePressEvent(self, event):
        if self.pan_mode and event.button() == Qt.MouseButton.LeftButton:
            self.last_pan_pos = event.globalPosition().toPoint()
            self.panning = True
            self.setCursor(Qt.CursorShape.ClosedHandCursor)
        else:
            super().mousePressEvent(event)

    def mouseMoveEvent(self, event):
        if self.panning and self.last_pan_pos:
            current_pos = event.globalPosition().toPoint()
            delta = current_pos - self.last_pan_pos
            self.last_pan_pos = current_pos
            
            h_bar = self.scroll_area.horizontalScrollBar()
            v_bar = self.scroll_area.verticalScrollBar()
            
            h_bar.setValue(h_bar.value() - delta.x())
            v_bar.setValue(v_bar.value() - delta.y())
        else:
            super().mouseMoveEvent(event)

    def mouseReleaseEvent(self, event):
        if self.panning and event.button() == Qt.MouseButton.LeftButton:
            self.panning = False
            self.last_pan_pos = None
            self.setCursor(Qt.CursorShape.OpenHandCursor)
        else:
            super().mouseReleaseEvent(event)


# --- CLASS BARU UNTUK WADAH TAB ---
class PdfTab(QWidget):
    """Wadah untuk satu dokumen PDF, lengkap dengan viewer dan statusnya."""
    def __init__(self, file_path, main_window):
        super().__init__(main_window)
        self.main_window = main_window
        self.file_path = file_path
        
        self.pdf = None
        self.current_page = 0
        self.zoom_factor = 1.0
        self.search_results = []
        self.current_search_index = -1
        
        # Buat UI untuk tab ini
        self.pdf_view_scroll_area = QScrollArea()
        self.pdf_view_scroll_area.setWidgetResizable(True)
        
        # Gunakan PdfViewLabel yang sudah kita punya
        self.pdf_view_label = PdfViewLabel(self.pdf_view_scroll_area)
        self.pdf_view_label.setAlignment(Qt.AlignmentFlag.AlignCenter)
        self.pdf_view_scroll_area.setWidget(self.pdf_view_label)
        
        # Atur layout untuk widget PdfTab ini
        layout = QVBoxLayout(self)
        layout.setContentsMargins(0, 0, 0, 0)
        layout.addWidget(self.pdf_view_scroll_area)
        
        # Terapkan mode pan dari tombol utama
        if self.main_window.pan_action.isChecked():
            self.pdf_view_label.setPanMode(True)

    def load_pdf(self, password=None):
        """Membuka dokumen PDF untuk tab ini."""
        try:
            self.pdf = pdfium.PdfDocument(self.file_path, password=password)
            self.go_to_page(0) # Tampilkan halaman pertama
            return True, None # Sukses, tidak perlu password
            
        except pdfium.PdfiumError as e:
            if "Password required" in str(e):
                return False, "password_required"
            if "Incorrect password" in str(e):
                return False, "incorrect_password"
            else:
                QMessageBox.critical(self.main_window, "Error", f"Failed to open PDF: {e}")
                return False, "error"
        except Exception as e:
            QMessageBox.critical(self.main_window, "Error", f"Failed to open PDF file: {e}")
            return False, "error"

    def close_pdf(self):
        """Menutup dokumen PDF di tab ini."""
        if self.pdf:
            self.pdf.close()
            self.pdf = None
            self.pdf_view_label.clear()

    def display_page(self, page_num):
        """Merender dan menampilkan halaman PDF tertentu."""
        if not self.pdf or not (0 <= page_num < len(self.pdf)):
            return

        self.current_page = page_num
        page = self.pdf[page_num]
        
        pil_image = page.render(scale=self.zoom_factor).to_pil()
        
        image = QImage(pil_image.tobytes(), pil_image.width, pil_image.height, pil_image.width * 3, QImage.Format.Format_RGB888)
        pixmap = QPixmap.fromImage(image)
        
        if self.search_results:
            self.draw_search_highlights(page, pixmap)
        
        self.pdf_view_label.setPixmap(pixmap)
        
        # Update UI utama
        self.main_window.update_status_bar()
        self.main_window.update_ui_state()

    def go_to_page(self, page_num):
        """Pindah ke halaman tertentu dan update thumbnail."""
        if self.pdf and 0 <= page_num < len(self.pdf):
            self.display_page(page_num)
            
            # Update highlight di thumbnail pane
            for i in range(self.main_window.thumb_layout.count()):
                widget = self.main_window.thumb_layout.itemAt(i).widget()
                if isinstance(widget, ThumbnailLabel):
                    if widget.page_num == page_num:
                        widget.setStyleSheet("border: 2px solid #0078d4; padding: 2px; margin: 5px;")
                    else:
                        widget.setStyleSheet("")

    def set_zoom(self, value):
        """Mengatur zoom dari nilai slider (persentase)."""
        self.zoom_factor = value / 100.0
        self.display_page(self.current_page)
        
    def set_zoom_factor(self, factor):
        """Mengatur zoom dari nilai float (misal 1.0) - berguna untuk session."""
        self.zoom_factor = factor
        self.display_page(self.current_page)

    def convert_pdf_rect_to_pixmap(self, page, rect):
        """Konversi koordinat PDF ke koordinat QPixmap."""
        l, b, r, t = rect
        page_width, page_height = page.get_size()
        
        pixmap_x = l * self.zoom_factor
        pixmap_y = (page_height - t) * self.zoom_factor
        pixmap_w = (r - l) * self.zoom_factor
        pixmap_h = (t - b) * self.zoom_factor
        
        return QRectF(pixmap_x, pixmap_y, pixmap_w, pixmap_h)

    def draw_search_highlights(self, page, pixmap):
        """Menggambar highlight hasil pencarian di atas pixmap."""
        painter = QPainter(pixmap)
        
        highlight_color = QColor(255, 255, 0, 100)
        active_highlight_color = QColor(255, 165, 0, 150)
        
        painter.setPen(Qt.PenStyle.NoPen)
        
        for i, (p_num, rect) in enumerate(self.search_results):
            if p_num == self.current_page:
                q_rect = self.convert_pdf_rect_to_pixmap(page, rect)
                
                if i == self.current_search_index:
                    painter.setBrush(active_highlight_color)
                else:
                    painter.setBrush(highlight_color)
                    
                painter.drawRect(q_rect)
        
        painter.end()
# --- AKHIR CLASS PDFTAB ---


class MacanReader(QMainWindow):
    def __init__(self, file_to_open=None):
        super().__init__()
        
        self.windows = []
        # Properti yang dulu di sini, sekarang pindah ke PdfTab
        
        self.settings = QSettings("MacanAngkasa", "MacanReader")

        self.themes_action_group = QActionGroup(self)
        self.themes_action_group.setExclusive(True)

        self.restore_window_state() # Pulihkan ukuran window SEBELUM init_ui
        
        self.init_ui()
        self.update_recent_files_menu()
        
        self.load_settings() # Muat tema

        # Coba restore session, jika gagal, baru buka file_to_open
        if not self.restore_session():
            if file_to_open:
                self.open_pdf_in_new_tab(file_to_open)
            else:
                self.add_placeholder_tab() # Tambahkan tab kosong jika tidak ada session

    def add_placeholder_tab(self):
        """Menambahkan tab 'Welcome' jika tidak ada PDF yang dibuka."""
        placeholder = QWidget()
        layout = QVBoxLayout(placeholder)
        layout.setAlignment(Qt.AlignmentFlag.AlignCenter)
        
        label = QLabel("Welcome to Macan Reader\n\nOpen file (Ctrl+O) or drag-and-drop PDF file here.\n\nMore info please visit: https://github.com/danx123/macan-reader/releases")
        label.setAlignment(Qt.AlignmentFlag.AlignCenter)
        font = label.font()
        font.setPointSize(14)
        label.setFont(font)
        
        layout.addWidget(label)
        
        index = self.tabs.addTab(placeholder, "Welcome")
        self.tabs.setCurrentIndex(index)

    def init_ui(self):
        self.setWindowTitle("Macan Reader")
        icon_path = "macan_reader.ico"
        if hasattr(sys, "_MEIPASS"):
            icon_path = os.path.join(sys._MEIPASS, icon_path)
        if os.path.exists(icon_path):
            self.setWindowIcon(QIcon(icon_path))
        
        screen_size = QApplication.primaryScreen().geometry()
        self.resize(int(screen_size.width() * 0.7), int(screen_size.height() * 0.8))

        main_widget = QWidget()
        main_layout = QVBoxLayout(main_widget)
        main_layout.setContentsMargins(0, 0, 0, 0)
        main_layout.setSpacing(0)

        self.search_widget = self.create_search_widget()
        main_layout.addWidget(self.search_widget)
        
        self.splitter = QSplitter(Qt.Orientation.Horizontal)
        # Jangan tambahkan splitter ke main_layout dulu
        
        self.setCentralWidget(main_widget)
        
        # --- Wrapper widget untuk thumbnail pane ---
        self.thumbnail_pane_widget = QWidget()
        thumbnail_pane_layout = QVBoxLayout(self.thumbnail_pane_widget)
        thumbnail_pane_layout.setContentsMargins(0, 0, 0, 0)
        thumbnail_pane_layout.setSpacing(0)
        
        thumb_title_bar = QWidget()
        thumb_title_bar.setObjectName("thumbnailTitleBar")
        thumb_title_bar_layout = QHBoxLayout(thumb_title_bar)
        thumb_title_bar_layout.setContentsMargins(5, 2, 2, 2)
        thumb_title_bar_layout.addWidget(QLabel("Thumbnails"))
        thumb_title_bar_layout.addStretch()
        
        self.close_thumb_btn = QPushButton("") 
        self.close_thumb_btn.setFlat(True)
        self.close_thumb_btn.setFixedSize(20, 20)
        self.close_thumb_btn.setToolTip("Close Thumbnail Pane")
        self.close_thumb_btn.clicked.connect(lambda: self.toggle_thumbnail_pane(False))
        thumb_title_bar_layout.addWidget(self.close_thumb_btn)
        
        thumbnail_pane_layout.addWidget(thumb_title_bar)
        # --- Akhir Penambahan ---

        self.thumb_scroll_area = QScrollArea()
        self.thumb_scroll_area.setWidgetResizable(True)
        self.thumb_container = QWidget()
        self.thumb_container.setObjectName("thumbnailContainer")
        self.thumb_layout = QVBoxLayout(self.thumb_container)
        self.thumb_layout.setAlignment(Qt.AlignmentFlag.AlignTop)
        self.thumb_scroll_area.setWidget(self.thumb_container)
        self.thumb_scroll_area.setMinimumWidth(150)
        
        thumbnail_pane_layout.addWidget(self.thumb_scroll_area)
        self.splitter.addWidget(self.thumbnail_pane_widget) # Tambahkan wrapper ke splitter
        
        # --- GANTI PDF VIEWER DENGAN TAB WIDGET ---
        self.tabs = QTabWidget()
        self.tabs.setTabsClosable(True)
        self.tabs.setMovable(True)
        self.tabs.tabCloseRequested.connect(self.close_tab)
        self.tabs.currentChanged.connect(self.update_on_tab_change)
        
        self.splitter.addWidget(self.tabs) # Tambahkan tab widget ke splitter
        # --- AKHIR PERUBAHAN ---
        
        main_layout.addWidget(self.splitter) # Tambahkan splitter ke main_layout
        
        self.splitter.setSizes([150, 800]) # Default sizes

        self.create_actions()
        self.create_menus()
        self.create_toolbar()
        self.create_status_bar()
        self.update_ui_state()
        
        self.setAcceptDrops(True)

    def current_tab(self):
        """Mengembalikan instance PdfTab yang sedang aktif."""
        widget = self.tabs.currentWidget()
        if isinstance(widget, PdfTab):
            return widget
        return None
    
    def new_window(self):
        """Membuka instance window MacanReader yang baru."""
        # Membuat objek window baru
        # file_to_open=None agar window baru mulai dengan Welcome Tab
        new_win = MacanReader(file_to_open=None) 
        
        # Simpan referensinya agar tetap ada
        self.windows.append(new_win)
        
        # Tampilkan window baru
        new_win.show()

    def create_search_widget(self):
        widget = QWidget()
        widget.setObjectName("searchWidget")
        layout = QHBoxLayout(widget)
        layout.setContentsMargins(5, 5, 5, 5)
        
        layout.addWidget(QLabel("Find:"))
        self.search_input = QLineEdit()
        self.search_input.returnPressed.connect(self.find_next)
        self.search_input.textChanged.connect(self.clear_search)
        layout.addWidget(self.search_input)
        
        find_next_btn = QPushButton("Next")
        find_next_btn.clicked.connect(self.find_next)
        layout.addWidget(find_next_btn)
        
        find_prev_btn = QPushButton("Previous")
        find_prev_btn.clicked.connect(self.find_prev)
        layout.addWidget(find_prev_btn)
        
        self.match_case_check = QCheckBox("Match Case")
        layout.addWidget(self.match_case_check)
        
        layout.addStretch()
        
        self.close_search_btn = QPushButton("")
        self.close_search_btn.setFlat(True)
        self.close_search_btn.setFixedSize(24, 24)
        self.close_search_btn.setToolTip("Close Search Bar")
        self.close_search_btn.clicked.connect(self.toggle_search_bar)
        layout.addWidget(self.close_search_btn)
        
        widget.setVisible(False)
        return widget

    def create_actions(self):
        default_icon_color = THEME_ICON_COLORS["light"]
        
        self.open_action = QAction(create_svg_icon(SVG_ICONS["open"], default_icon_color), "&Open...", self)
        self.open_action.triggered.connect(self.open_file_dialog)
        self.open_action.setShortcut("Ctrl+O")

        self.save_as_action = QAction(create_svg_icon(SVG_ICONS["save"], default_icon_color), "&Save As...", self)
        self.save_as_action.triggered.connect(self.save_as_file)

        self.export_png_action = QAction("Export as PNG...", self)
        self.export_png_action.triggered.connect(self.export_to_png)
        
        self.export_docx_action = QAction("Export as DOCX...", self)
        self.export_docx_action.triggered.connect(self.export_to_docx)

        self.print_action = QAction(create_svg_icon(SVG_ICONS["print"], default_icon_color), "&Print...", self)
        self.print_action.triggered.connect(self.print_file)
        self.print_action.setShortcut("Ctrl+P")

        self.close_action = QAction("&Close", self)
        self.close_action.triggered.connect(self.close_pdf) # Sekarang akan menutup tab
        self.close_action.setShortcut("Ctrl+W")
        
        self.clear_recent_action = QAction("Clear Recent Files", self)
        self.clear_recent_action.triggered.connect(self.clear_recent_files)

        self.exit_action = QAction("&Exit", self)
        self.exit_action.triggered.connect(self.close)
        self.exit_action.setShortcut("Ctrl+Q")
        
        self.find_action = QAction(create_svg_icon(SVG_ICONS["search"], default_icon_color), "&Find...", self)
        self.find_action.triggered.connect(self.toggle_search_bar)
        self.find_action.setShortcut("Ctrl+F")

        self.jump_to_page_action = QAction(create_svg_icon(SVG_ICONS["jump"], default_icon_color), "&Jump to Page...", self)
        self.jump_to_page_action.triggered.connect(self.jump_to_page_dialog)
        self.jump_to_page_action.setShortcut("Ctrl+G")

        self.new_window_action = QAction("New Window", self)
        self.new_window_action.triggered.connect(self.new_window)
        self.new_window_action.setShortcut("Ctrl+Shift+N")

        self.toggle_thumbnails_action = QAction("Show &Thumbnail Pane", self)
        self.toggle_thumbnails_action.setCheckable(True)
        self.toggle_thumbnails_action.setChecked(True)
        self.toggle_thumbnails_action.triggered.connect(self.toggle_thumbnail_pane)

        self.help_content_action = QAction("&Help Content", self)
        self.help_content_action.triggered.connect(self.show_help_content)
        self.help_content_action.setShortcut("F1")

        # --- BARIS BARU ---
        # Aksi untuk mendaftarkan sebagai default viewer
        self.set_default_action = QAction("Set as Default PDF Viewer...", self)
        self.set_default_action.triggered.connect(self.set_as_default_pdf_viewer)
        # --- AKHIR BARIS BARU ---

        self.about_action = QAction("&About Macan Reader", self)
        self.about_action.triggered.connect(self.show_about_dialog)
        
        self.prev_page_action = QAction(create_svg_icon(SVG_ICONS["prev-page"], default_icon_color), "Previous Page", self)
        self.prev_page_action.triggered.connect(lambda: self.go_to_page(self.current_tab().current_page - 1 if self.current_tab() else 0))
        self.prev_page_action.setShortcut("Left")

        self.next_page_action = QAction(create_svg_icon(SVG_ICONS["next-page"], default_icon_color), "Next Page", self)
        self.next_page_action.triggered.connect(lambda: self.go_to_page(self.current_tab().current_page + 1 if self.current_tab() else 0))
        self.next_page_action.setShortcut("Right")

        self.zoom_out_action = QAction(create_svg_icon(SVG_ICONS["zoom-out"], default_icon_color), "Zoom Out", self)
        self.zoom_out_action.triggered.connect(lambda: self.zoom_slider.setValue(self.zoom_slider.value() - 10))
        self.zoom_out_action.setShortcut("Ctrl+-")

        self.zoom_in_action = QAction(create_svg_icon(SVG_ICONS["zoom-in"], default_icon_color), "Zoom In", self)
        self.zoom_in_action.triggered.connect(lambda: self.zoom_slider.setValue(self.zoom_slider.value() + 10))
        self.zoom_in_action.setShortcut("Ctrl+=")
        
        self.zoom_reset_action = QAction(create_svg_icon(SVG_ICONS["zoom-reset"], default_icon_color), "Reset Zoom (100%)", self)
        self.zoom_reset_action.triggered.connect(self.reset_zoom)
        self.zoom_reset_action.setShortcut("Ctrl+0")

        self.pan_action = QAction(create_svg_icon(SVG_ICONS["hand"], default_icon_color), "Pan Tool (Drag to move page)", self)
        self.pan_action.setCheckable(True)
        self.pan_action.triggered.connect(self.toggle_pan_mode)
        
        # --- ACTION BARU UNTUK FULLSCREEN ---
        self.fullscreen_action = QAction("Fullscreen", self)
        self.fullscreen_action.setShortcut("F11")
        self.fullscreen_action.setCheckable(True)
        self.fullscreen_action.toggled.connect(self.toggle_fullscreen)
        # --- AKHIR ACTION BARU ---
        
        self.light_theme_action = QAction("Light", self, checkable=True, triggered=lambda: self.set_theme("light"))
        self.dark_theme_action = QAction("Dark", self, checkable=True, triggered=lambda: self.set_theme("dark"))
        self.dark_blue_theme_action = QAction("Dark Blue", self, checkable=True, triggered=lambda: self.set_theme("dark-blue"))
        self.neon_blue_theme_action = QAction("Neon Blue", self, checkable=True, triggered=lambda: self.set_theme("neon-blue"))
        self.soft_pink_theme_action = QAction("Soft Pink", self, checkable=True, triggered=lambda: self.set_theme("soft-pink"))
        
        self.themes_action_group.addAction(self.light_theme_action)
        self.themes_action_group.addAction(self.dark_theme_action)
        self.themes_action_group.addAction(self.dark_blue_theme_action)
        self.themes_action_group.addAction(self.neon_blue_theme_action)
        self.themes_action_group.addAction(self.soft_pink_theme_action)

    def create_menus(self):
        menu_bar = self.menuBar()
        
        file_menu = menu_bar.addMenu("&File")
        file_menu.addAction(self.open_action)
        file_menu.addAction(self.save_as_action)
        
        export_menu = file_menu.addMenu("Export As")
        export_menu.addAction(self.export_png_action)
        export_menu.addAction(self.export_docx_action)
        
        file_menu.addSeparator()
        file_menu.addAction(self.print_action)
        file_menu.addSeparator()
        
        self.recent_files_menu = file_menu.addMenu("Recent Files")
        
        file_menu.addAction(self.close_action)
        file_menu.addSeparator()
        file_menu.addAction(self.exit_action)
        
        edit_menu = menu_bar.addMenu("&Edit")
        edit_menu.addAction(self.find_action)
        edit_menu.addAction(self.jump_to_page_action)
        
        window_menu = menu_bar.addMenu("&Window")
        window_menu.addAction(self.new_window_action)
        window_menu.addAction(self.toggle_thumbnails_action)
        window_menu.addAction(self.fullscreen_action) # <-- TAMBAHKAN FULLSCREEN
        
        themes_menu = menu_bar.addMenu("&Themes")
        themes_menu.addAction(self.light_theme_action)
        themes_menu.addAction(self.dark_theme_action)
        themes_menu.addAction(self.dark_blue_theme_action)
        themes_menu.addAction(self.neon_blue_theme_action)
        themes_menu.addAction(self.soft_pink_theme_action)
        
        help_menu = menu_bar.addMenu("&Help")
        help_menu.addAction(self.help_content_action)
        help_menu.addSeparator()

        # --- BARIS BARU ---
        # Tambahkan aksi kita di sini
        help_menu.addAction(self.set_default_action)
        help_menu.addSeparator()
        # --- AKHIR BARIS BARU ---

        help_menu.addAction(self.about_action)

    def create_toolbar(self):
        toolbar = QToolBar("Main Toolbar")
        self.addToolBar(toolbar)
        
        toolbar.addAction(self.open_action)
        toolbar.addAction(self.print_action)
        toolbar.addAction(self.save_as_action)
        toolbar.addSeparator()

        toolbar.addAction(self.prev_page_action)
        toolbar.addAction(self.next_page_action)
        toolbar.addAction(self.jump_to_page_action)
        
        toolbar.addSeparator()

        toolbar.addAction(self.zoom_out_action)
        toolbar.addAction(self.zoom_in_action)
        toolbar.addAction(self.zoom_reset_action)
        
        toolbar.addSeparator()
        
        toolbar.addAction(self.pan_action)

    def create_status_bar(self):
        self.status_bar = QStatusBar()
        self.setStatusBar(self.status_bar)

        self.file_info_label = QLabel("No file opened")
        self.status_bar.addWidget(self.file_info_label, 1) # '1' artinya stretch

        # --- TOMBOL NAVIGASI ---
        self.status_prev_page_btn = QPushButton()
        self.status_prev_page_btn.setIcon(self.prev_page_action.icon())
        self.status_prev_page_btn.setToolTip("Previous Page (Left Arrow)")
        self.status_prev_page_btn.clicked.connect(self.prev_page_action.trigger)
        self.status_prev_page_btn.setFixedSize(28, 28)
        self.status_bar.addPermanentWidget(self.status_prev_page_btn)

        self.page_info_label = QLabel()
        self.status_bar.addPermanentWidget(self.page_info_label)

        self.status_next_page_btn = QPushButton()
        self.status_next_page_btn.setIcon(self.next_page_action.icon())
        self.status_next_page_btn.setToolTip("Next Page (Right Arrow)")
        self.status_next_page_btn.clicked.connect(self.next_page_action.trigger)
        self.status_next_page_btn.setFixedSize(28, 28)
        self.status_bar.addPermanentWidget(self.status_next_page_btn)
        # --- AKHIR TOMBOL NAVIGASI ---

        # --- TOMBOL RESET ZOOM ---
        self.status_zoom_reset_btn = QPushButton()
        self.status_zoom_reset_btn.setIcon(self.zoom_reset_action.icon())
        self.status_zoom_reset_btn.setToolTip("Reset Zoom (Ctrl+0)")
        self.status_zoom_reset_btn.clicked.connect(self.reset_zoom)
        self.status_zoom_reset_btn.setFixedSize(28, 28)
        self.status_bar.addPermanentWidget(self.status_zoom_reset_btn)
        # --- AKHIR TOMBOL RESET ZOOM ---

        self.zoom_slider = QSlider(Qt.Orientation.Horizontal)
        self.zoom_slider.setRange(25, 500)
        self.zoom_slider.setValue(100)
        self.zoom_slider.setFixedWidth(150)
        self.zoom_slider.valueChanged.connect(self.set_zoom)
        self.status_bar.addPermanentWidget(self.zoom_slider)

        self.zoom_label = QLabel(" 100% ")
        self.status_bar.addPermanentWidget(self.zoom_label)

        # --- TOMBOL FULLSCREEN ---
        self.status_fullscreen_btn = QPushButton()
        # Ikon di-set di update_icons_for_theme
        self.status_fullscreen_btn.setToolTip("Toggle Fullscreen (F11)")
        self.status_fullscreen_btn.setCheckable(True)
        self.status_fullscreen_btn.toggled.connect(self.toggle_fullscreen)
        self.status_fullscreen_btn.setFixedSize(28, 28)
        self.status_bar.addPermanentWidget(self.status_fullscreen_btn)
        # --- AKHIR TOMBOL FULLSCREEN ---

    def reset_zoom(self):
        self.zoom_slider.setValue(100)

    def load_settings(self):
        saved_theme = self.settings.value("theme", "light")
        
        if saved_theme == "dark":
            self.dark_theme_action.setChecked(True)
        elif saved_theme == "dark-blue":
            self.dark_blue_theme_action.setChecked(True)
        elif saved_theme == "neon-blue":
            self.neon_blue_theme_action.setChecked(True)
        elif saved_theme == "soft-pink":
            self.soft_pink_theme_action.setChecked(True)
        else:
            self.light_theme_action.setChecked(True)
            
        self.set_theme(saved_theme)

    def set_theme(self, theme_name):
        style = THEME_STYLESHEETS.get(theme_name, "")
        QApplication.instance().setStyleSheet(style)
        self.settings.setValue("theme", theme_name)
        
        icon_color = THEME_ICON_COLORS.get(theme_name, "black")
        self.update_icons_for_theme(icon_color)
        
    def update_icons_for_theme(self, icon_color):
        self.open_action.setIcon(create_svg_icon(SVG_ICONS["open"], icon_color))
        self.print_action.setIcon(create_svg_icon(SVG_ICONS["print"], icon_color))
        self.save_as_action.setIcon(create_svg_icon(SVG_ICONS["save"], icon_color))
        self.find_action.setIcon(create_svg_icon(SVG_ICONS["search"], icon_color))
        
        self.prev_page_action.setIcon(create_svg_icon(SVG_ICONS["prev-page"], icon_color))
        self.next_page_action.setIcon(create_svg_icon(SVG_ICONS["next-page"], icon_color))
        self.jump_to_page_action.setIcon(create_svg_icon(SVG_ICONS["jump"], icon_color))
        self.zoom_in_action.setIcon(create_svg_icon(SVG_ICONS["zoom-in"], icon_color))
        self.zoom_out_action.setIcon(create_svg_icon(SVG_ICONS["zoom-out"], icon_color))
        self.zoom_reset_action.setIcon(create_svg_icon(SVG_ICONS["zoom-reset"], icon_color))
        self.pan_action.setIcon(create_svg_icon(SVG_ICONS["hand"], icon_color))
        
        # Update ikon tombol status bar juga
        if hasattr(self, 'status_prev_page_btn'):
            self.status_prev_page_btn.setIcon(self.prev_page_action.icon())
            self.status_next_page_btn.setIcon(self.next_page_action.icon())
            self.status_zoom_reset_btn.setIcon(self.zoom_reset_action.icon())
            
            # Update ikon fullscreen
            if self.isFullScreen():
                self.status_fullscreen_btn.setIcon(create_svg_icon(SVG_ICONS["fullscreen-exit"], icon_color))
            else:
                self.status_fullscreen_btn.setIcon(create_svg_icon(SVG_ICONS["fullscreen-enter"], icon_color))
        
        close_icon_svg = """<svg xmlns="http://www.w3.org/2000/svg" width="16" height="16" viewBox="0 0 24 24" fill="none" stroke="currentColor" stroke-width="2" stroke-linecap="round" stroke-linejoin="round"><line x1="18" y1="6" x2="6" y2="18"></line><line x1="6" y1="6" x2="18" y2="18"></line></svg>"""
        self.close_search_btn.setIcon(create_svg_icon(close_icon_svg, icon_color))
        self.close_thumb_btn.setIcon(create_svg_icon(close_icon_svg, icon_color))

    def toggle_pan_mode(self, checked):
        # Terapkan ke semua tab yang terbuka
        for i in range(self.tabs.count()):
            widget = self.tabs.widget(i)
            if isinstance(widget, PdfTab):
                widget.pdf_view_label.setPanMode(checked)

    def toggle_thumbnail_pane(self, checked):
        self.thumbnail_pane_widget.setVisible(checked)
        if self.toggle_thumbnails_action.isChecked() != checked:
            self.toggle_thumbnails_action.setChecked(checked)

    def toggle_search_bar(self):
        is_visible = self.search_widget.isVisible()
        self.search_widget.setVisible(not is_visible)
        if not is_visible:
            self.search_input.setFocus()
        else:
            self.clear_search()

    def open_file_dialog(self):
        # Modifikasi untuk bisa multi-select
        file_paths, _ = QFileDialog.getOpenFileNames(self, "Open PDF File(s)", "", "PDF Files (*.pdf);;All Files (*)")
        if file_paths:
            for file_path in file_paths:
                self.open_pdf_in_new_tab(file_path)
            
    def open_pdf_in_new_tab(self, file_path):
        """Membuka file PDF di tab baru."""
        
        # Cek apakah file sudah terbuka
        for i in range(self.tabs.count()):
            widget = self.tabs.widget(i)
            if isinstance(widget, PdfTab) and widget.file_path == file_path:
                self.tabs.setCurrentIndex(i)
                return

        # Buat tab baru
        new_tab = PdfTab(file_path, self)
        
        # Coba load PDF
        success, status = new_tab.load_pdf()
        
        if status == "password_required":
            password, ok = QInputDialog.getText(self, "Password Required",
                                                f"PDF '{os.path.basename(file_path)}' terkunci. Masukkan password:",
                                                QLineEdit.Password)
            if ok and password:
                success, status = new_tab.load_pdf(password=password)
            else:
                new_tab.deleteLater() # Hapus tab jika user cancel
                return # Batal
        
        if not success:
            if status == "incorrect_password":
                QMessageBox.critical(self, "Error", "Password salah.")
            # Error lain sudah ditangani di dalam load_pdf()
            new_tab.deleteLater()
            return
            
        # Jika berhasil, tambahkan tab ke QTabWidget
        
        # Hapus tab "Welcome" jika ada
        if self.tabs.count() == 1 and not isinstance(self.tabs.widget(0), PdfTab):
            self.tabs.removeTab(0)
            
        index = self.tabs.addTab(new_tab, os.path.basename(file_path))
        self.tabs.setCurrentIndex(index)
        self.tabs.setTabToolTip(index, file_path) # Tambahkan tooltip
        
        self.add_to_recent_files(file_path)
        
        # update_on_tab_change akan dipanggil otomatis,
        # yang akan memanggil populate_thumbnails dan update UI

    def close_tab(self, index):
        """Slot untuk menutup tab saat tombol 'x' diklik."""
        widget = self.tabs.widget(index)
        if isinstance(widget, PdfTab):
            widget.close_pdf() # Tutup file pdfium
            widget.deleteLater() # Hapus widget
        
        self.tabs.removeTab(index)
        
        # Jika tidak ada tab tersisa, tambahkan placeholder
        if self.tabs.count() == 0:
            self.add_placeholder_tab()

    def close_pdf(self):
        """Menutup tab yang aktif saat ini."""
        current_index = self.tabs.currentIndex()
        if current_index != -1:
            self.close_tab(current_index)
            
    def update_on_tab_change(self, index):
        """Dipanggil setiap kali tab aktif berubah."""
        self.clear_search() # Hapus highlight pencarian
        
        tab = self.current_tab()
        if tab:
            self.populate_thumbnails() # Muat thumbnail untuk PDF di tab ini
            self.go_to_page(tab.current_page) # Update highlight thumbnail
            
            # Update zoom slider ke nilai zoom tab ini
            zoom_value = int(tab.zoom_factor * 100)
            self.zoom_slider.blockSignals(True)
            self.zoom_slider.setValue(zoom_value)
            self.zoom_slider.blockSignals(False)
            self.zoom_label.setText(f" {zoom_value}% ")
            
            self.setWindowTitle(f"Macan Reader - {os.path.basename(tab.file_path)}")
        else:
            self.populate_thumbnails() # Kosongkan thumbnail
            self.setWindowTitle("Macan Reader")

        self.update_status_bar()
        self.update_ui_state()

    def populate_thumbnails(self):
        while self.thumb_layout.count():
            child = self.thumb_layout.takeAt(0)
            if child.widget():
                child.widget().deleteLater()

        tab = self.current_tab()
        if not tab or not tab.pdf:
            return

        for i in range(len(tab.pdf)):
            page = tab.pdf[i]
            thumbnail = page.render(scale=0.15).to_pil()
            
            image = QImage(thumbnail.tobytes(), thumbnail.width, thumbnail.height, thumbnail.width * 3, QImage.Format.Format_RGB888)
            pixmap = QPixmap.fromImage(image)
            
            # --- Menggambar nomor halaman di atas pixmap ---
            painter = QPainter(pixmap)
            font = painter.font()
            font_size = max(10, int(pixmap.height() / 12))
            font.setPointSize(font_size)
            font.setBold(True)
            painter.setFont(font)
            
            text = str(i + 1)
            metrics = QFontMetrics(font)
            text_rect = metrics.boundingRect(text)
            
            bg_height = text_rect.height() + 4
            bg_width = text_rect.width() + 8
            bg_x = (pixmap.width() - bg_width) / 2
            bg_y = pixmap.height() - bg_height - 2
            
            bg_rect = QRectF(bg_x, bg_y, bg_width, bg_height)
            
            painter.setBrush(QColor(240, 240, 240, 200)) 
            painter.setPen(QColor(100, 100, 100, 150)) 
            painter.drawRoundedRect(bg_rect, 5, 5)
            
            painter.setPen(QColor(0, 0, 0)) 
            painter.drawText(bg_rect, Qt.AlignmentFlag.AlignCenter, text)
            painter.end()
            # --- AKHIR PENAMBAHAN ---
            
            thumb_label = ThumbnailLabel(i, self.thumb_container)
            thumb_label.setPixmap(pixmap)
            
            thumb_label.page_clicked.connect(self.go_to_page_from_thumb) # Slot baru
            thumb_label.export_requested.connect(self.export_single_page)
            
            self.thumb_layout.addWidget(thumb_label)
            
    def go_to_page_from_thumb(self, page_num):
        """Slot untuk klik thumbnail."""
        self.go_to_page(page_num)

    def go_to_page(self, page_num):
        """Pindah ke halaman di tab aktif."""
        tab = self.current_tab()
        if tab:
            tab.go_to_page(page_num)

    def jump_to_page_dialog(self):
        tab = self.current_tab()
        if not tab or not tab.pdf:
            return
            
        total_pages = len(tab.pdf)
        
        page_num, ok = QInputDialog.getInt(self, "Jump to Page",
                                           f"Enter page number (1 - {total_pages}):",
                                           tab.current_page + 1, 
                                           1,                      
                                           total_pages)            
        
        if ok and page_num:
            self.go_to_page(page_num - 1)

    def set_zoom(self, value):
        self.zoom_label.setText(f" {value}% ")
        tab = self.current_tab()
        if tab:
            tab.set_zoom(value) 

    # --- Bagian Logika Pencarian (Sudah di-refaktor) ---
    def clear_search(self):
        tab = self.current_tab()
        if tab:
            tab.search_results = []
            tab.current_search_index = -1
            tab.display_page(tab.current_page) 
        self.status_bar.clearMessage()

    def run_search(self):
        text = self.search_input.text()
        tab = self.current_tab()
        
        if not text or not tab or not tab.pdf:
            self.clear_search()
            return
            
        tab.search_results = []
        tab.current_search_index = -1
        
        match_case_enabled = self.match_case_check.isChecked()
        
        for i in range(len(tab.pdf)):
            page = tab.pdf[i]
            try:
                textpage = page.get_textpage()
            except Exception as e:
                print(f"Could not get textpage for page {i}: {e}")
                continue
                
            searcher = textpage.search(text, match_case=match_case_enabled)
            
            while True:
                rect = searcher.get_next() 
                if rect is None:
                    break
                
                if isinstance(rect, (list, tuple)) and len(rect) == 4:
                    tab.search_results.append((i, rect)) # Simpan ke tab
                else:
                    print(f"Warning: Skipping non-standard search rect on page {i}: {rect}")
                    pass
                
        self.status_bar.showMessage(f"Found {len(tab.search_results)} result(s).")
        
    def find_next(self):
        text = self.search_input.text()
        if not text: return
        
        tab = self.current_tab()
        if not tab: return
        
        if not tab.search_results:
            self.run_search()
            
        if not tab.search_results:
            return

        tab.current_search_index = (tab.current_search_index + 1) % len(tab.search_results)
        self.highlight_search_result()

    def find_prev(self):
        text = self.search_input.text()
        if not text: return

        tab = self.current_tab()
        if not tab: return

        if not tab.search_results:
            self.run_search()

        if not tab.search_results:
            return

        tab.current_search_index = (tab.current_search_index - 1 + len(tab.search_results)) % len(tab.search_results)
        self.highlight_search_result()

    def highlight_search_result(self):
        tab = self.current_tab()
        if not tab or not tab.search_results:
            return
            
        page_num, rect = tab.search_results[tab.current_search_index]
        
        if page_num != tab.current_page:
            self.go_to_page(page_num)
        else:
            tab.display_page(page_num) # Render ulang
            
        page = tab.pdf[page_num]
        q_rect = tab.convert_pdf_rect_to_pixmap(page, rect)
        
        h_bar = tab.pdf_view_scroll_area.horizontalScrollBar()
        v_bar = tab.pdf_view_scroll_area.verticalScrollBar()
        
        target_x = int(q_rect.center().x() - tab.pdf_view_scroll_area.viewport().width() / 2)
        target_y = int(q_rect.center().y() - tab.pdf_view_scroll_area.viewport().height() / 2)
        
        target_x = max(0, min(target_x, h_bar.maximum()))
        target_y = max(0, min(target_y, v_bar.maximum()))
        
        h_bar.setValue(target_x)
        v_bar.setValue(target_y)

    # --- Akhir Bagian Logika Pencarian ---

    def save_as_file(self):
        tab = self.current_tab()
        if not tab or not tab.pdf:
            return
        
        save_path, _ = QFileDialog.getSaveFileName(self, "Save PDF As", "", "PDF Files (*.pdf)")
        if save_path:
            try:
                tab.pdf.save(save_path)
                QMessageBox.information(self, "Success", f"File saved successfully to {save_path}")
            except Exception as e:
                QMessageBox.critical(self, "Error", f"Could not save file: {e}")

    # --- Bagian Ekspor (Sudah di-refaktor) ---

    def export_to_png(self):
        tab = self.current_tab()
        if not tab or not tab.pdf:
            return
        
        base_path, _ = QFileDialog.getSaveFileName(self, "Export Pages as PNG", 
                                                    f"{os.path.splitext(os.path.basename(tab.file_path))[0]}.png",
                                                    "PNG Files (*.png)")
        if not base_path:
            return
        
        base_name = os.path.splitext(base_path)[0]
        num_pages = len(tab.pdf)
        pad_width = len(str(num_pages))
        
        try:
            for i in range(num_pages):
                page = tab.pdf[i]
                img = page.render(scale=300/72.0).to_pil()
                file_name = f"{base_name}_page_{str(i+1).zfill(pad_width)}.png"
                img.save(file_name)
                
            QMessageBox.information(self, "Export Successful", 
                                    f"Exported {num_pages} pages to PNG files starting with '{os.path.basename(base_name)}'.")
        except Exception as e:
            QMessageBox.critical(self, "Export Error", f"Failed to export to PNG: {e}")

    def export_to_docx(self):
        tab = self.current_tab()
        if not tab or not tab.pdf:
            return
        
        save_path, _ = QFileDialog.getSaveFileName(self, "Export as DOCX", 
                                                   f"{os.path.splitext(os.path.basename(tab.file_path))[0]}.docx",
                                                    "Word Documents (*.docx)")
        if not save_path:
            return
        
        try:
            import docx
            from docx.shared import Inches
        except ImportError:
            QMessageBox.critical(self, "Missing Dependency",
                                 "Exporting to DOCX requires the 'python-docx' library.\n"
                                 "Please install it using: pip install python-docx")
            return

        try:
            doc = docx.Document()
            
            for i in range(len(tab.pdf)):
                page = tab.pdf[i]
                img = page.render(scale=150/72.0).to_pil()
                
                img_stream = io.BytesIO()
                img.save(img_stream, format='PNG')
                img_stream.seek(0)
                
                doc.add_picture(img_stream, width=Inches(6.5))
                
                if i < len(tab.pdf) - 1:
                    doc.add_page_break()
                        
            doc.save(save_path)
            QMessageBox.information(self, "Export Successful",
                                    f"Successfully exported PDF pages as images to {save_path}.")
        except Exception as e:
            QMessageBox.critical(self, "Export Error", f"Failed to export to DOCX: {e}")

    def export_single_page(self, page_num, format_type):
        tab = self.current_tab()
        if not tab or not tab.pdf or not (0 <= page_num < len(tab.pdf)):
            return
        
        base_name = os.path.splitext(os.path.basename(tab.file_path))[0]
        suggested_name = f"{base_name}_page_{page_num + 1}.{format_type}"
        
        if format_type == "png":
            self.export_page_as_png(page_num, suggested_name)
        elif format_type == "docx":
            self.export_page_as_docx(page_num, suggested_name)
        elif format_type == "pdf":
            self.export_page_as_pdf(page_num, suggested_name)

    def export_page_as_png(self, page_num, suggested_name):
        tab = self.current_tab()
        if not tab: return
        
        save_path, _ = QFileDialog.getSaveFileName(self, "Export Page as PNG", 
                                                   suggested_name, "PNG Files (*.png)")
        if not save_path:
            return
        
        try:
            page = tab.pdf[page_num]
            img = page.render(scale=300/72.0).to_pil() # 300 DPI
            img.save(save_path)
            QMessageBox.information(self, "Export Successful", f"Page {page_num + 1} saved to {save_path}")
        except Exception as e:
            QMessageBox.critical(self, "Export Error", f"Failed to export page to PNG: {e}")

    def export_page_as_docx(self, page_num, suggested_name):
        tab = self.current_tab()
        if not tab: return
        
        save_path, _ = QFileDialog.getSaveFileName(self, "Export Page as DOCX",
                                                   suggested_name, "Word Documents (*.docx)")
        if not save_path:
            return
        
        try:
            import docx
            from docx.shared import Inches
        except ImportError:
            QMessageBox.critical(self, "Missing Dependency",
                                 "Exporting to DOCX requires the 'python-docx' library.\n"
                                 "Please install it using: pip install python-docx")
            return
            
        try:
            doc = docx.Document()
            page = tab.pdf[page_num]
            img = page.render(scale=150/72.0).to_pil() # 150 DPI
            
            img_stream = io.BytesIO()
            img.save(img_stream, format='PNG')
            img_stream.seek(0)
            
            doc.add_picture(img_stream, width=Inches(6.5))
            doc.save(save_path)
            QMessageBox.information(self, "Export Successful", f"Page {page_num + 1} saved to {save_path}")
        except Exception as e:
            QMessageBox.critical(self, "Export Error", f"Failed to export page to DOCX: {e}")

    def export_page_as_pdf(self, page_num, suggested_name):
        tab = self.current_tab()
        if not tab: return
        
        save_path, _ = QFileDialog.getSaveFileName(self, "Export Page as PDF",
                                                   suggested_name, "PDF Files (*.pdf)")
        if not save_path:
            return
            
        try:
            new_pdf = pdfium.PdfDocument.new()
            new_pdf.import_pages(tab.pdf, [page_num]) # Gunakan tab.pdf
            new_pdf.save(save_path)
            new_pdf.close()
            QMessageBox.information(self, "Export Successful", f"Page {page_num + 1} saved to {save_path}")
        except Exception as e:
            QMessageBox.critical(self, "Export Error", f"Failed to export page to PDF: {e}")

    # --- Akhir Bagian Ekspor ---

    def print_file(self):
        tab = self.current_tab()
        if not tab or not tab.pdf:
            return

        printer = QPrinter(QPrinter.PrinterMode.HighResolution)
        preview_dialog = QPrintPreviewDialog(printer, self)
        preview_dialog.paintRequested.connect(self.handle_paint_request)
        preview_dialog.exec()

    def handle_paint_request(self, printer):
        tab = self.current_tab()
        if not tab:
            return
            
        painter = QPainter()
        painter.begin(printer)
        
        print_range_setting = printer.printRange()
        from_page = printer.fromPage()
        to_page = printer.toPage()
        
        pages_to_print = list(range(len(tab.pdf)))

        if print_range_setting == QPrinter.PrintRange.PageRange:
            pages_to_print = list(range(from_page - 1, min(to_page, len(tab.pdf))))
        elif print_range_setting == QPrinter.PrintRange.CurrentPage:
            pages_to_print = [tab.current_page]

        for i, page_idx in enumerate(pages_to_print):
            if page_idx >= len(tab.pdf):
                continue
                
            page = tab.pdf[page_idx]
            
            dpi = printer.resolution()
            if dpi == 0:
                dpi = 300
            scale = dpi / 72.0
            
            img = page.render(scale=scale).to_pil()
            
            q_img = QImage(img.tobytes(), img.width, img.height, img.width * 3, QImage.Format.Format_RGB888)
            
            page_rect = printer.pageRect(QPrinter.Unit.DevicePixel)
            img_pixmap = QPixmap.fromImage(q_img)
            
            scaled_pixmap = img_pixmap.scaled(page_rect.size().toSize(), Qt.AspectRatioMode.KeepAspectRatio, Qt.TransformationMode.SmoothTransformation)
            
            x = (page_rect.width() - scaled_pixmap.width()) / 2
            y = (page_rect.height() - scaled_pixmap.height()) / 2
            
            painter.drawPixmap(int(x), int(y), scaled_pixmap)
            
            if i < len(pages_to_print) - 1:
                printer.newPage()
        painter.end()

    def show_help_content(self):
        help_text = """
        <h2>Macan Reader - Help</h2>
        
        <h3>Navigation</h3>
        <ul>
        <li><b>Open File (Ctrl+O):</b> Open one or more PDF files. The file will open in a new tab.</li>
        <li><b>Close Tab (Ctrl+W):</b> Closes the active PDF tab.</li>
        <li><b>Switch Tabs:</b> Click on the tab title.</li>
        </ul>

        <h3>Reading PDFs</h3>
        <ul>
        <li><b>Next Page:</b> Click the 'Next' arrow button (in the toolbar or status bar) or press the <b>Right Arrow</b>.</li>
        <li><b>Previous Page:</b> Click the 'Previous' arrow button (in the toolbar or status bar) or press the <b>Left Arrow</b>.</li>
        <li><b>Go to Page:</b> Click the thumbnail in the left panel.</li>
        <li><b>Jump to Page (Ctrl+G):</b> Use the 'Jump' key (<b>#</b>) to insert a page number.</li>
        </ul>

        <h3>Zoom & Panning</h3>
        <ul>
        <li><b>Zoom In/Out:</b> Use the <b>+</b> and <b>-</b> keys, the slider in the status bar, or <b>Ctrl+</b>/<b>Ctrl-</b>.</li>
        <li><b>Reset Zoom:</b> Click the 'target' button (in the toolbar or status bar) or press <b>Ctrl+0</b> to reset to 100%.</li>
        <li><b>Pan:</b> Click the 'Hand' icon to enter Pan mode. Click and drag the page to pan.</li>
        </ul>

        <h3>Other Features</h3>
        <ul>
        <li><b>Search (Ctrl+F):</b> Type in the 'Find' bar above. Click <b>Next</b>/<b>Previous</b> to navigate.</li>
        <li><b>Fullscreen (F11):</b> Use the button in the status bar or <b>Window > Fullscreen</b>.</li>
        <li><b>Thumbnail:</b> Right-click a thumbnail to export that page.</li>
        <li><b>Theme:</b> Change the appearance of the application via the <b>Themes</b> menu.</li>
        <li><b>Session:</b> The application will remember open tabs when closed and reopened.</li>
        </ul>
        """
        
        help_dialog = QMessageBox(self)
        help_dialog.setWindowTitle("Macan Reader Help")
        help_dialog.setIcon(QMessageBox.Icon.Question)
        
        help_dialog.setText(" ") 
        
        scroll_area = QScrollArea(help_dialog)
        scroll_area.setWidgetResizable(True)
        scroll_area.setMinimumSize(500, 300)
        
        help_label = QLabel(help_text, scroll_area)
        help_label.setTextFormat(Qt.TextFormat.RichText)
        help_label.setWordWrap(True)
        help_label.setAlignment(Qt.AlignmentFlag.AlignTop)
        
        scroll_area.setWidget(help_label)
        
        layout = help_dialog.layout()
        
        old_widget = layout.itemAtPosition(0, 1).widget()
        if old_widget:
            old_widget.deleteLater()
            
        layout.addWidget(scroll_area, 0, 1, 1, 1)
        
        help_dialog.exec()

    def show_about_dialog(self):
        about_text = """
        <h2>Macan Reader</h2>
        <p>A professional PDF reader built with PySide6 and PyPDFium2.</p>
        <p>Fast, efficient, multi-tab, and reliable.</p>
        <p>v4.0.0 (Tabbed Edition)</p>
        <br>
        <p>© 2025 - Danx Exodus - Macan Angkasa</p>
        """
        QMessageBox.about(self, "About Macan Reader", about_text)

    # --- FUNGSI BARU DI SINI ---
    def set_as_default_pdf_viewer(self):
        """Mencoba mendaftarkan Macan Reader sebagai default PDF viewer (Hanya Windows)."""
        
        # 1. Cek apakah ini Windows
        if sys.platform != "win32":
            QMessageBox.information(self, "Not Supported",
                                    "This feature is only available on Windows.")
            return

        # 2. Import modul khusus Windows
        try:
            import winreg
            import ctypes
        except ImportError:
            QMessageBox.critical(self, "Error",
                                 "Failed to import required Windows modules (winreg, ctypes).")
            return

        # 3. Dapatkan path ke .exe
        # PENTING: Ini hanya berfungsi jika sudah di-bundle!
        exe_path = sys.executable
        if not exe_path.lower().endswith(".exe"):
            resp = QMessageBox.warning(self, "Warning",
                                       "This function is intended to be run from a compiled .exe file.\n"
                                       f"Registering the current path ({exe_path}) will likely not work correctly.\n\n"
                                       "Do you want to continue anyway?",
                                       QMessageBox.StandardButton.Yes | QMessageBox.StandardButton.No)
            if resp == QMessageBox.StandardButton.No:
                return
        
        prog_id = "MacanReader.pdf"
        command = f'"{exe_path}" "%1"'
        app_name = "Macan Reader"

        try:
            # Tulis ke HKEY_CURRENT_USER (tidak perlu admin)
            # 4. Daftarkan ProgID
            key = winreg.CreateKey(winreg.HKEY_CURRENT_USER, fr"Software\Classes\{prog_id}")
            winreg.SetValue(key, None, winreg.REG_SZ, "PDF Document (Macan Reader)")
            winreg.CloseKey(key)

            # 5. Atur ikon default
            key = winreg.CreateKey(winreg.HKEY_CURRENT_USER, fr"Software\Classes\{prog_id}\DefaultIcon")
            winreg.SetValue(key, None, winreg.REG_SZ, f'"{exe_path}",0')
            winreg.CloseKey(key)

            # 6. Atur perintah 'open'
            key = winreg.CreateKey(winreg.HKEY_CURRENT_USER, fr"Software\Classes\{prog_id}\shell\open\command")
            winreg.SetValue(key, None, winreg.REG_SZ, command)
            winreg.CloseKey(key)
            
            # 7. Kaitkan ekstensi .pdf dengan ProgID
            key = winreg.CreateKey(winreg.HKEY_CURRENT_USER, r"Software\Classes\.pdf")
            winreg.SetValue(key, None, winreg.REG_SZ, prog_id)
            winreg.CloseKey(key)

            # 8. Daftarkan aplikasi di "RegisteredApplications" (penting untuk "Default Programs")
            key = winreg.CreateKey(winreg.HKEY_CURRENT_USER, r"Software\RegisteredApplications")
            winreg.SetValueEx(key, app_name, 0, winreg.REG_SZ, r"Software\MacanReader\Capabilities")
            winreg.CloseKey(key)

            # 9. Tentukan kapabilitas aplikasi
            key = winreg.CreateKey(winreg.HKEY_CURRENT_USER, r"Software\MacanReader\Capabilities")
            winreg.SetValueEx(key, "ApplicationName", 0, winreg.REG_SZ, app_name)
            winreg.SetValueEx(key, "ApplicationDescription", 0, winreg.REG_SZ, "Macan Reader PDF Viewer")
            winreg.CloseKey(key)

            key = winreg.CreateKey(winreg.HKEY_CURRENT_USER, r"Software\MacanReader\Capabilities\FileAssociations")
            winreg.SetValueEx(key, ".pdf", 0, winreg.REG_SZ, prog_id)
            winreg.CloseKey(key)

            # 10. Beri tahu Windows bahwa ada perubahan asosiasi file
            SHCNE_ASSOCCHANGED = 0x08000000
            SHCNF_IDLIST = 0x0000
            ctypes.windll.shell32.SHChangeNotify(SHCNE_ASSOCCHANGED, SHCNF_IDLIST, None, None)

            QMessageBox.information(self, "Success",
                                    f"{app_name} has been registered as a PDF handler.\n\n"
                                    "On modern Windows, you may still need to manually select it as the default app:\n"
                                    "Go to: Settings > Apps > Default apps")

        except PermissionError:
            QMessageBox.warning(self, "Permission Denied",
                                "Could not write to registry. Try running as administrator.")
        except Exception as e:
            QMessageBox.critical(self, "Registry Error",
                                 f"An error occurred while writing to the registry:\n{e}")
    # --- AKHIR FUNGSI BARU ---

    def add_to_recent_files(self, file_path):
        recent_files = self.settings.value("recentFiles", [])
        if file_path in recent_files:
            recent_files.remove(file_path)
        recent_files.insert(0, file_path)
        del recent_files[10:]
        self.settings.setValue("recentFiles", recent_files)
        self.update_recent_files_menu()
        
    def update_recent_files_menu(self):
        self.recent_files_menu.clear()
        recent_files = self.settings.value("recentFiles", [])
        
        if recent_files:
            for file_path in recent_files:
                action = QAction(os.path.basename(file_path), self)
                action.setData(file_path)
                action.triggered.connect(self.open_recent_file)
                self.recent_files_menu.addAction(action)
            self.recent_files_menu.addSeparator()
            self.recent_files_menu.addAction(self.clear_recent_action)
        else:
            no_recent_action = QAction("No Recent Files", self)
            no_recent_action.setEnabled(False)
            self.recent_files_menu.addAction(no_recent_action)

    def clear_recent_files(self):
        self.settings.setValue("recentFiles", [])
        self.update_recent_files_menu()

    def open_recent_file(self):
        action = self.sender()
        if action:
            file_path = action.data()
            if os.path.exists(file_path):
                self.open_pdf_in_new_tab(file_path) # Buka di tab baru
            else:
                QMessageBox.warning(self, "File Not Found", f"The file '{file_path}' could not be found.")
                recent_files = self.settings.value("recentFiles", [])
                if file_path in recent_files:
                    recent_files.remove(file_path)
                    self.settings.setValue("recentFiles", recent_files)
                    self.update_recent_files_menu()

    def update_ui_state(self):
        tab = self.current_tab()
        is_pdf_loaded = (tab is not None and tab.pdf is not None)
        
        self.save_as_action.setEnabled(is_pdf_loaded)
        self.print_action.setEnabled(is_pdf_loaded)
        self.close_action.setEnabled(is_pdf_loaded)
        self.export_png_action.setEnabled(is_pdf_loaded)
        self.export_docx_action.setEnabled(is_pdf_loaded)
        
        self.find_action.setEnabled(is_pdf_loaded)
        self.jump_to_page_action.setEnabled(is_pdf_loaded)
        self.pan_action.setEnabled(is_pdf_loaded)
        if not is_pdf_loaded and self.pan_action.isChecked():
            self.pan_action.setChecked(False)
            self.toggle_pan_mode(False)
        
        self.zoom_in_action.setEnabled(is_pdf_loaded)
        self.zoom_out_action.setEnabled(is_pdf_loaded)
        self.zoom_reset_action.setEnabled(is_pdf_loaded)
        self.zoom_slider.setEnabled(is_pdf_loaded)
        
        # Update tombol status bar
        if hasattr(self, 'status_prev_page_btn'):
            self.status_zoom_reset_btn.setEnabled(is_pdf_loaded)
        
        if is_pdf_loaded:
            self.prev_page_action.setEnabled(tab.current_page > 0)
            self.next_page_action.setEnabled(tab.current_page < len(tab.pdf) - 1)
            # Update tombol status bar juga
            if hasattr(self, 'status_prev_page_btn'):
                self.status_prev_page_btn.setEnabled(tab.current_page > 0)
                self.status_next_page_btn.setEnabled(tab.current_page < len(tab.pdf) - 1)
        else:
            self.prev_page_action.setEnabled(False)
            self.next_page_action.setEnabled(False)
            if hasattr(self, 'status_prev_page_btn'):
                self.status_prev_page_btn.setEnabled(False)
                self.status_next_page_btn.setEnabled(False)

    def update_status_bar(self):
        tab = self.current_tab()
        if tab:
            self.file_info_label.setText(os.path.basename(tab.file_path))
            self.page_info_label.setText(f" Page: {tab.current_page + 1} of {len(tab.pdf)} ")
        else:
            self.file_info_label.setText("No file opened")
            self.page_info_label.setText("")

    def dragEnterEvent(self, event):
        if event.mimeData().hasUrls():
            for url in event.mimeData().urls():
                if url.isLocalFile() and url.toLocalFile().lower().endswith('.pdf'):
                    event.acceptProposedAction()
                    return
        event.ignore()

    def dropEvent(self, event):
        if event.mimeData().hasUrls():
            event.acceptProposedAction()
            for url in event.mimeData().urls():
                if url.isLocalFile() and url.toLocalFile().lower().endswith('.pdf'):
                    file_path = url.toLocalFile()
                    self.open_pdf_in_new_tab(file_path)
        else:
            super().dropEvent(event)
            
    # --- FUNGSI BARU UNTUK FULLSCREEN ---
    def toggle_fullscreen(self, checked):
        """Mengaktifkan atau menonaktifkan mode fullscreen."""
        if checked:
            self.showFullScreen()
        else:
            self.showNormal()
        
        # Update tombol dan action (untuk sinkronisasi)
        self.status_fullscreen_btn.setChecked(checked)
        self.fullscreen_action.setChecked(checked)
        
        # Update ikon tombol
        current_theme = self.settings.value("theme", "light")
        icon_color = THEME_ICON_COLORS.get(current_theme, "black")
        if checked:
            self.status_fullscreen_btn.setIcon(create_svg_icon(SVG_ICONS["fullscreen-exit"], icon_color))
        else:
            self.status_fullscreen_btn.setIcon(create_svg_icon(SVG_ICONS["fullscreen-enter"], icon_color))

    def changeEvent(self, event):
        """Menangani event perubahan state window (termasuk exit F11 manual)."""
        if event.type() == event.Type.WindowStateChange:
            is_fullscreen = self.isFullScreen()
            # Sinkronkan tombol jika state berubah (misal user pencet Esc)
            if self.status_fullscreen_btn.isChecked() != is_fullscreen:
                self.toggle_fullscreen(is_fullscreen)
        
        super().changeEvent(event)

    # --- FUNGSI BARU UNTUK SESSION MANAGEMENT ---
    def save_window_state(self):
        """Menyimpan geometri dan state window."""
        self.settings.setValue("geometry", self.saveGeometry())
        self.settings.setValue("windowState", self.saveState())
        self.settings.setValue("splitterState", self.splitter.saveState())
    
    def restore_window_state(self):
        """Memulihkan geometri dan state window."""
        geometry = self.settings.value("geometry")
        if geometry:
            self.restoreGeometry(geometry)
        
        state = self.settings.value("windowState")
        if state:
            self.restoreState(state)
            
        splitterState = self.settings.value("splitterState")
        if splitterState:
            try:
                self.splitter.restoreState(splitterState)
            except Exception as e:
                print(f"Warning: Could not restore splitter state: {e}")

    def save_session(self):
        """Menyimpan semua tab yang terbuka ke settings."""
        tabs_data = []
        for i in range(self.tabs.count()):
            widget = self.tabs.widget(i)
            if isinstance(widget, PdfTab):
                tab_info = {
                    "file_path": widget.file_path,
                    "current_page": widget.current_page,
                    "zoom_factor": widget.zoom_factor
                }
                tabs_data.append(tab_info)
        
        if tabs_data:
            self.settings.setValue("session/tabs", tabs_data)
            self.settings.setValue("session/current_index", self.tabs.currentIndex())
        else:
            self.settings.remove("session") # Hapus session jika tidak ada tab

    def restore_session(self):
        """Memulihkan tab dari session sebelumnya."""
        tabs_data = self.settings.value("session/tabs", [], type=list)
        if not tabs_data:
            return False # Tidak ada session untuk dipulihkan

        for tab_info in tabs_data:
            file_path = tab_info.get("file_path")
            if file_path and os.path.exists(file_path):
                self.open_pdf_in_new_tab(file_path)
                
                new_tab_widget = self.current_tab() 
                if new_tab_widget and new_tab_widget.file_path == file_path:
                    
                    zoom_factor = tab_info.get("zoom_factor", 1.0)
                    new_tab_widget.set_zoom_factor(zoom_factor)
                    
                    page = tab_info.get("current_page", 0)
                    new_tab_widget.go_to_page(page)
            
        current_index = self.settings.value("session/current_index", 0, type=int)
        if current_index < self.tabs.count():
            self.tabs.setCurrentIndex(current_index)
        
        return self.tabs.count() > 0

    def closeEvent(self, event):
        self.save_session()
        self.save_window_state()
        
        # Tutup semua PDF secara manual sebelum keluar
        for i in range(self.tabs.count()):
            widget = self.tabs.widget(i)
            if isinstance(widget, PdfTab):
                widget.close_pdf()
        
        self.settings.sync()
        event.accept()

if __name__ == "__main__":
    app = QApplication(sys.argv)
    
    file_to_open = None
    if len(sys.argv) > 1:
        path = sys.argv[1]
        if os.path.exists(path) and path.lower().endswith('.pdf'):
            file_to_open = path
            
    # file_to_open diteruskan ke init,
    # tapi init akan memprioritaskan restore_session
    window = MacanReader(file_to_open)
    window.show()
    sys.exit(app.exec())