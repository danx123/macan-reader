import sys
import os
import pypdfium2 as pdfium
import io  # DITAMBAHKAN: Untuk ekspor ke memory buffer

from PySide6.QtWidgets import (
    QApplication, QMainWindow, QWidget, QVBoxLayout, QHBoxLayout, QLabel,
    QScrollArea, QSplitter, QFileDialog, QStatusBar, QSlider, QMessageBox,
    QToolBar, QPushButton, QLineEdit, QCheckBox, QMenu,
    QInputDialog  # DITAMBAHKAN: Untuk password dan jump to page
)
from PySide6.QtGui import (
    QPixmap, QImage, QIcon, QAction, QPainter, QScreen, QColor,
    QActionGroup,
    QFont, QFontMetrics # DITAMBAHKAN: Untuk menggambar teks di thumbnail
)
from PySide6.QtCore import (
    Qt, QSize, QSettings, QUrl, QByteArray, Signal, QRectF
)
from PySide6.QtSvg import QSvgRenderer
# DIUBAH: Menambahkan QPrintPreviewDialog dan menghapus QPrintDialog
from PySide6.QtPrintSupport import QPrinter, QPrintPreviewDialog

# --- Helper untuk menangani path saat di-bundle dengan PyInstaller ---
def resource_path(relative_path):
    """ Dapatkan path absolut ke resource, berfungsi untuk dev dan PyInstaller """
    try:
        base_path = sys._MEIPASS
    except Exception:
        base_path = os.path.abspath(".")
    return os.path.join(base_path, relative_path)

# --- Icon SVG dalam format XML (Embedded) ---
SVG_ICONS = {
    "open": """<svg xmlns="http://www.w3.org/2000/svg" width="24" height="24" viewBox="0 0 24 24" fill="none" stroke="currentColor" stroke-width="2" stroke-linecap="round" stroke-linejoin="round"><path d="M22 19a2 2 0 0 1-2 2H4a2 2 0 0 1-2-2V5a2 2 0 0 1 2-2h5l2 3h9a2 2 0 0 1 2 2z"></path></svg>""",
    "print": """<svg xmlns="http://www.w3.org/2000/svg" width="24" height="24" viewBox="0 0 24 24" fill="none" stroke="currentColor" stroke-width="2" stroke-linecap="round" stroke-linejoin="round"><polyline points="6 9 6 2 18 2 18 9"></polyline><path d="M6 18H4a2 2 0 0 1-2-2v-5a2 2 0 0 1 2-2h16a2 2 0 0 1 2 2v5a2 2 0 0 1-2 2h-2"></path><rect x="6" y="14" width="12" height="8"></rect></svg>""",
    "zoom-in": """<svg xmlns="http://www.w3.org/2000/svg" width="24" height="24" viewBox="0 0 24 24" fill="none" stroke="currentColor" stroke-width="2" stroke-linecap="round" stroke-linejoin="round"><circle cx="11" cy="11" r="8"></circle><line x1="21" y1="21" x2="16.65" y2="16.65"></line><line x1="11" y1="8" x2="11" y2="14"></line><line x1="8" y1="11" x2="14" y2="11"></line></svg>""",
    "zoom-out": """<svg xmlns="http://www.w3.org/2000/svg" width="24" height="24" viewBox="0 0 24 24" fill="none" stroke="currentColor" stroke-width="2" stroke-linecap="round" stroke-linejoin="round"><circle cx="11" cy="11" r="8"></circle><line x1="21" y1="21" x2="16.65" y2="16.65"></line><line x1="8" y1="11" x2="14" y2="11"></line></svg>""",
    "zoom-reset": """<svg xmlns="http://www.w3.org/2000/svg" width="24" height="24" viewBox="0 0 24 24" fill="none" stroke="currentColor" stroke-width="2" stroke-linecap="round" stroke-linejoin="round"><circle cx="12" cy="12" r="10"></circle><circle cx="12" cy="12" r="6"></circle><circle cx="12" cy="12" r="2"></circle></svg>""",
    "prev-page": """<svg xmlns="http://www.w3.org/2000/svg" width="24" height="24" viewBox="0 0 24 24" fill="none" stroke="currentColor" stroke-width="2" stroke-linecap="round" stroke-linejoin="round"><polyline points="15 18 9 12 15 6"></polyline></svg>""",
    "next-page": """<svg xmlns="http://www.w3.org/2000/svg" width="24" height="24" viewBox="0 0 24 24" fill="none" stroke="currentColor" stroke-width="2" stroke-linecap="round" stroke-linejoin="round"><polyline points="9 18 15 12 9 6"></polyline></svg>""",
    # DITAMBAHKAN: Ikon untuk Jump to Page (hashtag/pound symbol)
    "jump": """<svg xmlns="http://www.w3.org/2000/svg" width="24" height="24" viewBox="0 0 24 24" fill="none" stroke="currentColor" stroke-width="2" stroke-linecap="round" stroke-linejoin="round"><line x1="4" y1="9" x2="20" y2="9"></line><line x1="4" y1="15" x2="20" y2="15"></line><line x1="10" y1="3" x2="8" y2="21"></line><line x1="16" y1="3" x2="14" y2="21"></line></svg>""",
    "save": """<svg xmlns="http://www.w3.org/2000/svg" width="24" height="24" viewBox="0 0 24 24" fill="none" stroke="currentColor" stroke-width="2" stroke-linecap="round" stroke-linejoin="round"><path d="M19 21H5a2 2 0 0 1-2-2V5a2 2 0 0 1 2-2h11l5 5v11a2 2 0 0 1-2 2z"></path><polyline points="17 21 17 13 7 13 7 21"></polyline><polyline points="7 3 7 8 15 8"></polyline></svg>""",
    "search": """<svg xmlns="http://www.w3.org/2000/svg" width="24" height="24" viewBox="0 0 24 24" fill="none" stroke="currentColor" stroke-width="2" stroke-linecap="round" stroke-linejoin="round"><circle cx="11" cy="11" r="8"></circle><line x1="21" y1="21" x2="16.65" y2="16.65"></line></svg>""",
    "hand": """<svg xmlns="http://www.w3.org/2000/svg" width="24" height="24" viewBox="0 0 24 24" fill="none" stroke="currentColor" stroke-width="2" stroke-linecap="round" stroke-linejoin="round"><path d="M18 11V6a2 2 0 0 0-2-2v0a2 2 0 0 0-2 2v0"></path><path d="M14 10V4a2 2 0 0 0-2-2v0a2 2 0 0 0-2 2v2"></path><path d="M10 10.5V6a2 2 0 0 0-2-2v0a2 2 0 0 0-2 2v8"></path><path d="M18 8a2 2 0 0 1 2 2v10a2 2 0 0 1-2 2h-4a2 2 0 0 1-2-2v-4a2 2 0 0 1 2-2h2"></path></svg>"""
}

# DITAMBAHKAN: Kamus untuk warna ikon tema
THEME_ICON_COLORS = {
    "light": "black",
    "dark": "#E0E0E0",
    "dark-blue": "#E0E0E0",
    "neon-blue": "#E0E0E0",
    "soft-pink": "#333333"
}

# DITAMBAHKAN: Kamus untuk QSS Tema
THEME_STYLESHEETS = {
    "light": """
        QWidget {
            background-color: #FFFFFF;
            color: #000000;
        }
        QMainWindow, QToolBar, QStatusBar, QMenu, QMenuBar {
            background-color: #F0F0F0;
            color: #000000;
        }
        QMenuBar::item:selected, QMenu::item:selected {
            background-color: #0078D4;
            color: #FFFFFF;
        }
        QScrollArea {
            background-color: #E0E0E0;
            border: 1px solid #C0C0C0;
        }
        QPushButton {
            background-color: #E1E1E1;
            border: 1px solid #C0C0C0;
            padding: 5px;
        }
        QPushButton:hover {
            background-color: #E5E5E5;
        }
        QPushButton:pressed {
            background-color: #D0D0D0;
        }
        QLineEdit {
            background-color: #FFFFFF;
            border: 1px solid #C0C0C0;
        }
        QSlider::groove:horizontal {
            background: #C0C0C0;
            height: 8px;
            border-radius: 4px;
        }
        QSlider::handle:horizontal {
            background: #0078D4;
            width: 18px;
            height: 18px;
            border-radius: 9px;
            margin: -5px 0;
        }
        QSplitter::handle {
            background-color: #C0C0C0;
        }
        QWidget#searchWidget, QWidget#thumbnailTitleBar {
            background-color: #F0F0F0;
        }
        QWidget#thumbnailContainer {
            background-color: #E0E0E0;
        }
        ThumbnailLabel {
            border: 1px solid lightgray;
            padding: 2px;
            margin: 5px;
        }
        ThumbnailLabel:hover {
            border: 1px solid dodgerblue;
        }
    """,
    "dark": """
        QWidget {
            background-color: #3C3C3C;
            color: #E0E0E0;
        }
        QMainWindow, QToolBar, QStatusBar, QMenu, QMenuBar {
            background-color: #2D2D2D;
            color: #E0E0E0;
        }
        QMenuBar::item:selected, QMenu::item:selected {
            background-color: #0078D4;
            color: #FFFFFF;
        }
        QScrollArea {
            background-color: #333333;
            border: 1px solid #4A4A4A;
        }
        QPushButton {
            background-color: #4A4A4A;
            border: 1px solid #5A5A5A;
            padding: 5px;
        }
        QPushButton:hover {
            background-color: #5A5A5A;
        }
        QPushButton:pressed {
            background-color: #444444;
        }
        QLineEdit {
            background-color: #333333;
            border: 1px solid #5A5A5A;
        }
        QSlider::groove:horizontal {
            background: #4A4A4A;
            height: 8px;
            border-radius: 4px;
        }
        QSlider::handle:horizontal {
            background: #0078D4;
            width: 18px;
            height: 18px;
            border-radius: 9px;
            margin: -5px 0;
        }
        QSplitter::handle {
            background-color: #4A4A4A;
        }
        QWidget#searchWidget, QWidget#thumbnailTitleBar {
            background-color: #2D2D2D;
        }
        QWidget#thumbnailContainer {
            background-color: #333333;
        }
        ThumbnailLabel {
            border: 1px solid #5A5A5A;
            padding: 2px;
            margin: 5px;
        }
        ThumbnailLabel:hover {
            border: 1px solid #0078D4;
        }
    """,
    "dark-blue": """
        QWidget {
            background-color: #2D3A4F;
            color: #E0E0E0;
        }
        QMainWindow, QToolBar, QStatusBar, QMenu, QMenuBar {
            background-color: #222B3A;
            color: #E0E0E0;
        }
        QMenuBar::item:selected, QMenu::item:selected {
            background-color: #4A90E2;
            color: #FFFFFF;
        }
        QScrollArea {
            background-color: #283344;
            border: 1px solid #3A4A63;
        }
        QPushButton {
            background-color: #3A4A63;
            border: 1px solid #4A5B79;
            padding: 5px;
        }
        QPushButton:hover {
            background-color: #4A5B79;
        }
        QPushButton:pressed {
            background-color: #33425A;
        }
        QLineEdit {
            background-color: #283344;
            border: 1px solid #4A5B79;
        }
        QSlider::groove:horizontal {
            background: #3A4A63;
            height: 8px;
            border-radius: 4px;
        }
        QSlider::handle:horizontal {
            background: #4A90E2;
            width: 18px;
            height: 18px;
            border-radius: 9px;
            margin: -5px 0;
        }
        QSplitter::handle {
            background-color: #3A4A63;
        }
        QWidget#searchWidget, QWidget#thumbnailTitleBar {
            background-color: #222B3A;
        }
        QWidget#thumbnailContainer {
            background-color: #283344;
        }
        ThumbnailLabel {
            border: 1px solid #3A4A63;
            padding: 2px;
            margin: 5px;
        }
        ThumbnailLabel:hover {
            border: 1px solid #4A90E2;
        }
    """,
    "neon-blue": """
        QWidget {
            background-color: #0A0A1A;
            color: #00FFD1;
        }
        QMainWindow, QToolBar, QStatusBar, QMenu, QMenuBar {
            background-color: #0F0F2A;
            color: #00FFD1;
            border: 1px solid #00FFD1;
        }
        QMenuBar::item:selected, QMenu::item:selected {
            background-color: #00FFD1;
            color: #000000;
        }
        QScrollArea {
            background-color: #050510;
            border: 1px solid #00FFD1;
        }
        QPushButton {
            background-color: #1A1A3A;
            border: 1px solid #00FFD1;
            color: #00FFD1;
            padding: 5px;
        }
        QPushButton:hover {
            background-color: #2A2A4A;
        }
        QPushButton:pressed {
            background-color: #11112A;
        }
        QLineEdit {
            background-color: #050510;
            border: 1px solid #00FFD1;
            color: #00FFD1;
        }
        QSlider::groove:horizontal {
            background: #1A1A3A;
            height: 8px;
            border-radius: 4px;
        }
        QSlider::handle:horizontal {
            background: #00FFD1;
            width: 18px;
            height: 18px;
            border-radius: 9px;
            margin: -5px 0;
            border: 1px solid #00FFD1;
        }
        QSplitter::handle {
            background-color: #1A1A3A;
            border: 1px solid #00FFD1;
        }
        QWidget#searchWidget, QWidget#thumbnailTitleBar {
            background-color: #0F0F2A;
        }
        QWidget#thumbnailContainer {
            background-color: #050510;
        }
        ThumbnailLabel {
            border: 1px solid #00FFD1;
            padding: 2px;
            margin: 5px;
        }
        ThumbnailLabel:hover {
            border: 2px solid #00FFD1;
        }
    """,
    "soft-pink": """
        QWidget {
            background-color: #FFF0F5;
            color: #5D4037;
        }
        QMainWindow, QToolBar, QStatusBar, QMenu, QMenuBar {
            background-color: #FFDDEE;
            color: #5D4037;
        }
        QMenuBar::item:selected, QMenu::item:selected {
            background-color: #FF80AB;
            color: #FFFFFF;
        }
        QScrollArea {
            background-color: #FFF5F9;
            border: 1px solid #FFC0CB;
        }
        QPushButton {
            background-color: #FFC0CB;
            border: 1px solid #FF80AB;
            color: #5D4037;
            padding: 5px;
        }
        QPushButton:hover {
            background-color: #FFB0C1;
        }
        QPushButton:pressed {
            background-color: #FFA0B1;
        }
        QLineEdit {
            background-color: #FFF5F9;
            border: 1px solid #FFC0CB;
            color: #5D4037;
        }
        QSlider::groove:horizontal {
            background: #FFC0CB;
            height: 8px;
            border-radius: 4px;
        }
        QSlider::handle:horizontal {
            background: #FF80AB;
            width: 18px;
            height: 18px;
            border-radius: 9px;
            margin: -5px 0;
        }
        QSplitter::handle {
            background-color: #FFC0CB;
        }
        QWidget#searchWidget, QWidget#thumbnailTitleBar {
            background-color: #FFDDEE;
        }
        QWidget#thumbnailContainer {
            background-color: #FFF5F9;
        }
        ThumbnailLabel {
            border: 1px solid #FFC0CB;
            padding: 2px;
            margin: 5px;
        }
        ThumbnailLabel:hover {
            border: 1px solid #FF80AB;
        }
    """
}


def create_svg_icon(svg_xml, color="black"):
    """Membuat QIcon dari string XML SVG."""
    svg_xml = svg_xml.replace('stroke="currentColor"', f'stroke="{color}"')
    renderer = QSvgRenderer(QByteArray(svg_xml.encode('utf-8')))
    pixmap = QPixmap(QSize(24, 24))
    pixmap.fill(Qt.GlobalColor.transparent)
    painter = QPainter(pixmap)
    renderer.render(painter)
    painter.end()
    return QIcon(pixmap)

class ThumbnailLabel(QLabel):
    """Label kustom untuk thumbnail yang bisa diklik."""
    page_clicked = Signal(int)
    export_requested = Signal(int, str) # DITAMBAHKAN: Sinyal untuk ekspor

    def __init__(self, page_num, parent=None):
        super().__init__(parent)
        self.page_num = page_num
        self.setToolTip(f"Go to page {self.page_num + 1}")
        self.setCursor(Qt.CursorShape.PointingHandCursor)
        # DITAMBAHKAN: Aktifkan context menu
        self.setContextMenuPolicy(Qt.ContextMenuPolicy.CustomContextMenu)
        self.customContextMenuRequested.connect(self.show_context_menu)

    def mousePressEvent(self, event):
        if event.button() == Qt.MouseButton.LeftButton:
            self.page_clicked.emit(self.page_num)
        # DIUBAH: Jangan proses klik kanan di sini agar context menu bisa muncul
        elif event.button() != Qt.MouseButton.RightButton:
            super().mousePressEvent(event)

    # DITAMBAHKAN: Fungsi untuk menampilkan context menu
    def show_context_menu(self, pos):
        context_menu = QMenu(self)
        export_menu = context_menu.addMenu("Export Page As...")
        
        export_pdf_action = export_menu.addAction("PDF...")
        export_png_action = export_menu.addAction("PNG...")
        export_docx_action = export_menu.addAction("DOCX...")
        
        action = context_menu.exec(self.mapToGlobal(pos))
        
        if action == export_pdf_action:
            self.export_requested.emit(self.page_num, "pdf")
        elif action == export_png_action:
            self.export_requested.emit(self.page_num, "png")
        elif action == export_docx_action:
            self.export_requested.emit(self.page_num, "docx")

# DITAMBAHKAN: Kelas label kustom untuk panning
class PdfViewLabel(QLabel):
    """Label kustom untuk tampilan PDF yang mendukung panning."""
    def __init__(self, scroll_area, parent=None):
        super().__init__(parent)
        self.scroll_area = scroll_area
        self.pan_mode = False
        self.panning = False
        self.last_pan_pos = None
        self.setCursor(Qt.CursorShape.ArrowCursor)

    def setPanMode(self, enabled):
        """Mengaktifkan atau menonaktifkan mode pan."""
        self.pan_mode = enabled
        self.setCursor(Qt.CursorShape.OpenHandCursor if enabled else Qt.CursorShape.ArrowCursor)

    def mousePressEvent(self, event):
        if self.pan_mode and event.button() == Qt.MouseButton.LeftButton:
            self.last_pan_pos = event.globalPosition().toPoint()
            self.panning = True
            self.setCursor(Qt.CursorShape.ClosedHandCursor)
        else:
            super().mousePressEvent(event)

    def mouseMoveEvent(self, event):
        if self.panning and self.last_pan_pos:
            current_pos = event.globalPosition().toPoint()
            delta = current_pos - self.last_pan_pos
            self.last_pan_pos = current_pos
            
            h_bar = self.scroll_area.horizontalScrollBar()
            v_bar = self.scroll_area.verticalScrollBar()
            
            h_bar.setValue(h_bar.value() - delta.x())
            v_bar.setValue(v_bar.value() - delta.y())
        else:
            super().mouseMoveEvent(event)

    def mouseReleaseEvent(self, event):
        if self.panning and event.button() == Qt.MouseButton.LeftButton:
            self.panning = False
            self.last_pan_pos = None
            self.setCursor(Qt.CursorShape.OpenHandCursor)
        else:
            super().mouseReleaseEvent(event)

class MacanReader(QMainWindow):
    def __init__(self, file_to_open=None):
        super().__init__()
        
        self.pdf = None
        self.current_page = 0
        self.zoom_factor = 1.0
        self.current_file_path = None
        
        self.search_results = []
        self.current_search_index = -1
        
        self.settings = QSettings("MacanAngkasa", "MacanReader")

        self.themes_action_group = QActionGroup(self)
        self.themes_action_group.setExclusive(True)

        self.init_ui()
        self.update_recent_files_menu()
        
        self.load_settings()

        if file_to_open:
            self.open_pdf(file_to_open)

    def init_ui(self):
        self.setWindowTitle("Macan Reader")
        icon_path = "icon.ico"
        if hasattr(sys, "_MEIPASS"):
            icon_path = os.path.join(sys._MEIPASS, icon_path)
        if os.path.exists(icon_path):
            self.setWindowIcon(QIcon(icon_path))
        
        screen_size = QApplication.primaryScreen().geometry()
        self.resize(int(screen_size.width() * 0.7), int(screen_size.height() * 0.8))

        main_widget = QWidget()
        main_layout = QVBoxLayout(main_widget)
        main_layout.setContentsMargins(0, 0, 0, 0)
        main_layout.setSpacing(0)

        self.search_widget = self.create_search_widget()
        main_layout.addWidget(self.search_widget)
        
        self.splitter = QSplitter(Qt.Orientation.Horizontal)
        main_layout.addWidget(self.splitter)
        
        self.setCentralWidget(main_widget)
        
        # --- DITAMBAHKAN: Wrapper widget untuk thumbnail pane ---
        self.thumbnail_pane_widget = QWidget()
        thumbnail_pane_layout = QVBoxLayout(self.thumbnail_pane_widget)
        thumbnail_pane_layout.setContentsMargins(0, 0, 0, 0)
        thumbnail_pane_layout.setSpacing(0)
        
        thumb_title_bar = QWidget()
        thumb_title_bar.setObjectName("thumbnailTitleBar")
        thumb_title_bar_layout = QHBoxLayout(thumb_title_bar)
        thumb_title_bar_layout.setContentsMargins(5, 2, 2, 2)
        thumb_title_bar_layout.addWidget(QLabel("Thumbnails"))
        thumb_title_bar_layout.addStretch()
        
        self.close_thumb_btn = QPushButton("") # Icon di-set di update_icons
        self.close_thumb_btn.setFlat(True)
        self.close_thumb_btn.setFixedSize(20, 20)
        self.close_thumb_btn.setToolTip("Close Thumbnail Pane")
        self.close_thumb_btn.clicked.connect(lambda: self.toggle_thumbnail_pane(False))
        thumb_title_bar_layout.addWidget(self.close_thumb_btn)
        
        thumbnail_pane_layout.addWidget(thumb_title_bar)
        # --- Akhir Penambahan ---

        self.thumb_scroll_area = QScrollArea()
        self.thumb_scroll_area.setWidgetResizable(True)
        self.thumb_container = QWidget()
        self.thumb_container.setObjectName("thumbnailContainer")
        self.thumb_layout = QVBoxLayout(self.thumb_container)
        self.thumb_layout.setAlignment(Qt.AlignmentFlag.AlignTop)
        self.thumb_scroll_area.setWidget(self.thumb_container)
        self.thumb_scroll_area.setMinimumWidth(150)
        
        # --- DITAMBAHKAN: Tambahkan scroll area ke wrapper ---
        thumbnail_pane_layout.addWidget(self.thumb_scroll_area)
        self.splitter.addWidget(self.thumbnail_pane_widget) # Tambahkan wrapper ke splitter
        # --- Akhir Penambahan ---
        
        self.pdf_view_scroll_area = QScrollArea()
        self.pdf_view_scroll_area.setWidgetResizable(True)
        
        self.pdf_view_label = PdfViewLabel(self.pdf_view_scroll_area)
        self.pdf_view_label.setAlignment(Qt.AlignmentFlag.AlignCenter)
        self.pdf_view_scroll_area.setWidget(self.pdf_view_label)
        self.splitter.addWidget(self.pdf_view_scroll_area)
        
        self.splitter.setSizes([150, 800])

        self.create_actions()
        self.create_menus()
        self.create_toolbar()
        self.create_status_bar()
        self.update_ui_state()
        
        self.setAcceptDrops(True)

    def create_search_widget(self):
        widget = QWidget()
        widget.setObjectName("searchWidget")
        layout = QHBoxLayout(widget)
        layout.setContentsMargins(5, 5, 5, 5)
        
        layout.addWidget(QLabel("Find:"))
        self.search_input = QLineEdit()
        self.search_input.returnPressed.connect(self.find_next)
        self.search_input.textChanged.connect(self.clear_search)
        layout.addWidget(self.search_input)
        
        find_next_btn = QPushButton("Next")
        find_next_btn.clicked.connect(self.find_next)
        layout.addWidget(find_next_btn)
        
        find_prev_btn = QPushButton("Previous")
        find_prev_btn.clicked.connect(self.find_prev)
        layout.addWidget(find_prev_btn)
        
        self.match_case_check = QCheckBox("Match Case")
        layout.addWidget(self.match_case_check)
        
        layout.addStretch()
        
        self.close_search_btn = QPushButton("")
        self.close_search_btn.setFlat(True)
        self.close_search_btn.setFixedSize(24, 24)
        self.close_search_btn.setToolTip("Close Search Bar")
        self.close_search_btn.clicked.connect(self.toggle_search_bar)
        layout.addWidget(self.close_search_btn)
        
        widget.setVisible(False)
        return widget

    def create_actions(self):
        default_icon_color = THEME_ICON_COLORS["light"]
        
        self.open_action = QAction(create_svg_icon(SVG_ICONS["open"], default_icon_color), "&Open...", self)
        self.open_action.triggered.connect(self.open_file_dialog)
        self.open_action.setShortcut("Ctrl+O")

        self.save_as_action = QAction(create_svg_icon(SVG_ICONS["save"], default_icon_color), "&Save As...", self)
        self.save_as_action.triggered.connect(self.save_as_file)

        # DITAMBAHKAN: Aksi Ekspor
        self.export_png_action = QAction("Export as PNG...", self)
        self.export_png_action.triggered.connect(self.export_to_png)
        
        self.export_docx_action = QAction("Export as DOCX...", self)
        self.export_docx_action.triggered.connect(self.export_to_docx)

        self.print_action = QAction(create_svg_icon(SVG_ICONS["print"], default_icon_color), "&Print...", self)
        self.print_action.triggered.connect(self.print_file) # DIUBAH: Akan memanggil print preview
        self.print_action.setShortcut("Ctrl+P")

        self.close_action = QAction("&Close", self)
        self.close_action.triggered.connect(self.close_pdf)
        self.close_action.setShortcut("Ctrl+W")
        
        self.clear_recent_action = QAction("Clear Recent Files", self)
        self.clear_recent_action.triggered.connect(self.clear_recent_files)

        self.exit_action = QAction("&Exit", self)
        self.exit_action.triggered.connect(self.close)
        self.exit_action.setShortcut("Ctrl+Q")
        
        self.find_action = QAction(create_svg_icon(SVG_ICONS["search"], default_icon_color), "&Find...", self)
        self.find_action.triggered.connect(self.toggle_search_bar)
        self.find_action.setShortcut("Ctrl+F")

        # DITAMBAHKAN: Aksi Jump to Page
        self.jump_to_page_action = QAction(create_svg_icon(SVG_ICONS["jump"], default_icon_color), "&Jump to Page...", self)
        self.jump_to_page_action.triggered.connect(self.jump_to_page_dialog)
        self.jump_to_page_action.setShortcut("Ctrl+G")

        self.toggle_thumbnails_action = QAction("Show &Thumbnail Pane", self)
        self.toggle_thumbnails_action.setCheckable(True)
        self.toggle_thumbnails_action.setChecked(True)
        self.toggle_thumbnails_action.triggered.connect(self.toggle_thumbnail_pane)

        # DITAMBAHKAN: Aksi Help Content
        self.help_content_action = QAction("&Help Content", self)
        self.help_content_action.triggered.connect(self.show_help_content)
        self.help_content_action.setShortcut("F1")

        self.about_action = QAction("&About Macan Reader", self)
        self.about_action.triggered.connect(self.show_about_dialog)
        
        self.prev_page_action = QAction(create_svg_icon(SVG_ICONS["prev-page"], default_icon_color), "Previous Page", self)
        self.prev_page_action.triggered.connect(lambda: self.go_to_page(self.current_page - 1))
        self.prev_page_action.setShortcut("Left")

        self.next_page_action = QAction(create_svg_icon(SVG_ICONS["next-page"], default_icon_color), "Next Page", self)
        self.next_page_action.triggered.connect(lambda: self.go_to_page(self.current_page + 1))
        self.next_page_action.setShortcut("Right")

        self.zoom_out_action = QAction(create_svg_icon(SVG_ICONS["zoom-out"], default_icon_color), "Zoom Out", self)
        self.zoom_out_action.triggered.connect(lambda: self.zoom_slider.setValue(self.zoom_slider.value() - 10))
        self.zoom_out_action.setShortcut("Ctrl+-")

        self.zoom_in_action = QAction(create_svg_icon(SVG_ICONS["zoom-in"], default_icon_color), "Zoom In", self)
        self.zoom_in_action.triggered.connect(lambda: self.zoom_slider.setValue(self.zoom_slider.value() + 10))
        self.zoom_in_action.setShortcut("Ctrl+=")
        
        self.zoom_reset_action = QAction(create_svg_icon(SVG_ICONS["zoom-reset"], default_icon_color), "Reset Zoom (100%)", self)
        self.zoom_reset_action.triggered.connect(self.reset_zoom)
        self.zoom_reset_action.setShortcut("Ctrl+0")

        self.pan_action = QAction(create_svg_icon(SVG_ICONS["hand"], default_icon_color), "Pan Tool (Drag to move page)", self)
        self.pan_action.setCheckable(True)
        self.pan_action.triggered.connect(self.toggle_pan_mode)
        
        self.light_theme_action = QAction("Light", self, checkable=True, triggered=lambda: self.set_theme("light"))
        self.dark_theme_action = QAction("Dark", self, checkable=True, triggered=lambda: self.set_theme("dark"))
        self.dark_blue_theme_action = QAction("Dark Blue", self, checkable=True, triggered=lambda: self.set_theme("dark-blue"))
        self.neon_blue_theme_action = QAction("Neon Blue", self, checkable=True, triggered=lambda: self.set_theme("neon-blue"))
        self.soft_pink_theme_action = QAction("Soft Pink", self, checkable=True, triggered=lambda: self.set_theme("soft-pink"))
        
        self.themes_action_group.addAction(self.light_theme_action)
        self.themes_action_group.addAction(self.dark_theme_action)
        self.themes_action_group.addAction(self.dark_blue_theme_action)
        self.themes_action_group.addAction(self.neon_blue_theme_action)
        self.themes_action_group.addAction(self.soft_pink_theme_action)


    def create_menus(self):
        menu_bar = self.menuBar()
        
        file_menu = menu_bar.addMenu("&File")
        file_menu.addAction(self.open_action)
        file_menu.addAction(self.save_as_action)
        
        # --- DITAMBAHKAN: Menu Ekspor ---
        export_menu = file_menu.addMenu("Export As")
        export_menu.addAction(self.export_png_action)
        export_menu.addAction(self.export_docx_action)
        # --- Akhir Penambahan ---
        
        file_menu.addSeparator()
        file_menu.addAction(self.print_action)
        file_menu.addSeparator()
        
        self.recent_files_menu = file_menu.addMenu("Recent Files")
        
        file_menu.addAction(self.close_action)
        file_menu.addSeparator()
        file_menu.addAction(self.exit_action)
        
        edit_menu = menu_bar.addMenu("&Edit")
        edit_menu.addAction(self.find_action)
        edit_menu.addAction(self.jump_to_page_action) # DITAMBAHKAN: Jump to Page
        
        window_menu = menu_bar.addMenu("&Window")
        window_menu.addAction(self.toggle_thumbnails_action)
        
        themes_menu = menu_bar.addMenu("&Themes")
        themes_menu.addAction(self.light_theme_action)
        themes_menu.addAction(self.dark_theme_action)
        themes_menu.addAction(self.dark_blue_theme_action)
        themes_menu.addAction(self.neon_blue_theme_action)
        themes_menu.addAction(self.soft_pink_theme_action)
        
        help_menu = menu_bar.addMenu("&Help")
        # DITAMBAHKAN: Menu Help Content
        help_menu.addAction(self.help_content_action)
        help_menu.addSeparator()
        help_menu.addAction(self.about_action)

    def create_toolbar(self):
        toolbar = QToolBar("Main Toolbar")
        self.addToolBar(toolbar)
        
        toolbar.addAction(self.open_action)
        toolbar.addAction(self.print_action)
        toolbar.addAction(self.save_as_action)
        toolbar.addSeparator()

        toolbar.addAction(self.prev_page_action)
        toolbar.addAction(self.next_page_action)
        toolbar.addAction(self.jump_to_page_action) # DITAMBAHKAN: Jump to Page
        
        toolbar.addSeparator()

        toolbar.addAction(self.zoom_out_action)
        toolbar.addAction(self.zoom_in_action)
        toolbar.addAction(self.zoom_reset_action)
        
        toolbar.addSeparator()
        
        toolbar.addAction(self.pan_action)

    def create_status_bar(self):
        self.status_bar = QStatusBar()
        self.setStatusBar(self.status_bar)

        self.file_info_label = QLabel("No file opened")
        self.status_bar.addWidget(self.file_info_label, 1)

        self.page_info_label = QLabel()
        self.status_bar.addPermanentWidget(self.page_info_label)

        self.zoom_slider = QSlider(Qt.Orientation.Horizontal)
        self.zoom_slider.setRange(25, 500)
        self.zoom_slider.setValue(100)
        self.zoom_slider.setFixedWidth(150)
        self.zoom_slider.valueChanged.connect(self.set_zoom)
        self.status_bar.addPermanentWidget(self.zoom_slider)

        self.zoom_label = QLabel(" 100% ")
        self.status_bar.addPermanentWidget(self.zoom_label)

    def reset_zoom(self):
        self.zoom_slider.setValue(100)

    def load_settings(self):
        saved_theme = self.settings.value("theme", "light")
        
        if saved_theme == "dark":
            self.dark_theme_action.setChecked(True)
        elif saved_theme == "dark-blue":
            self.dark_blue_theme_action.setChecked(True)
        elif saved_theme == "neon-blue":
            self.neon_blue_theme_action.setChecked(True)
        elif saved_theme == "soft-pink":
            self.soft_pink_theme_action.setChecked(True)
        else:
            self.light_theme_action.setChecked(True)
            
        self.set_theme(saved_theme)

    def set_theme(self, theme_name):
        style = THEME_STYLESHEETS.get(theme_name, "")
        QApplication.instance().setStyleSheet(style)
        self.settings.setValue("theme", theme_name)
        
        icon_color = THEME_ICON_COLORS.get(theme_name, "black")
        self.update_icons_for_theme(icon_color)
        
    def update_icons_for_theme(self, icon_color):
        self.open_action.setIcon(create_svg_icon(SVG_ICONS["open"], icon_color))
        self.print_action.setIcon(create_svg_icon(SVG_ICONS["print"], icon_color))
        self.save_as_action.setIcon(create_svg_icon(SVG_ICONS["save"], icon_color))
        self.find_action.setIcon(create_svg_icon(SVG_ICONS["search"], icon_color))
        
        self.prev_page_action.setIcon(create_svg_icon(SVG_ICONS["prev-page"], icon_color))
        self.next_page_action.setIcon(create_svg_icon(SVG_ICONS["next-page"], icon_color))
        self.jump_to_page_action.setIcon(create_svg_icon(SVG_ICONS["jump"], icon_color)) # DITAMBAHKAN
        self.zoom_in_action.setIcon(create_svg_icon(SVG_ICONS["zoom-in"], icon_color))
        self.zoom_out_action.setIcon(create_svg_icon(SVG_ICONS["zoom-out"], icon_color))
        self.zoom_reset_action.setIcon(create_svg_icon(SVG_ICONS["zoom-reset"], icon_color))
        self.pan_action.setIcon(create_svg_icon(SVG_ICONS["hand"], icon_color))
        
        close_icon_svg = """<svg xmlns="http://www.w3.org/2000/svg" width="16" height="16" viewBox="0 0 24 24" fill="none" stroke="currentColor" stroke-width="2" stroke-linecap="round" stroke-linejoin="round"><line x1="18" y1="6" x2="6" y2="18"></line><line x1="6" y1="6" x2="18" y2="18"></line></svg>"""
        self.close_search_btn.setIcon(create_svg_icon(close_icon_svg, icon_color))
        # DITAMBAHKAN: Set ikon untuk tombol close thumbnail
        self.close_thumb_btn.setIcon(create_svg_icon(close_icon_svg, icon_color))

    def toggle_pan_mode(self, checked):
        self.pdf_view_label.setPanMode(checked)

    def toggle_thumbnail_pane(self, checked):
        # DIUBAH: Gunakan wrapper widget
        self.thumbnail_pane_widget.setVisible(checked)
        # Sinkronkan checkbox di menu
        if self.toggle_thumbnails_action.isChecked() != checked:
            self.toggle_thumbnails_action.setChecked(checked)

    def toggle_search_bar(self):
        is_visible = self.search_widget.isVisible()
        self.search_widget.setVisible(not is_visible)
        if not is_visible:
            self.search_input.setFocus()
        else:
            self.clear_search()

    def open_file_dialog(self):
        file_path, _ = QFileDialog.getOpenFileName(self, "Open PDF File", "", "PDF Files (*.pdf);;All Files (*)")
        if file_path:
            self.open_pdf(file_path)
            
    # --- FUNGSI open_pdf DIUBAH TOTAL UNTUK MENANGANI PASSWORD ---
    def open_pdf(self, file_path):
        try:
            if self.pdf:
                self.close_pdf()
                
            temp_pdf = None
            try:
                # Coba buka tanpa password dulu
                temp_pdf = pdfium.PdfDocument(file_path)
            except pdfium.PdfiumError as e:
                # Jika error adalah "Password required"
                if "Password required" in str(e):
                    # Minta password ke user
                    password, ok = QInputDialog.getText(self, "Password Required",
                                                        "This PDF is password protected. Please enter the password:",
                                                        QLineEdit.Password)
                    # Jika user memasukkan password dan klik OK
                    if ok and password:
                        try:
                            # Coba buka lagi dengan password
                            temp_pdf = pdfium.PdfDocument(file_path, password=password)
                        except pdfium.PdfiumError as e2:
                            # Jika password salah
                            if "Incorrect password" in str(e2):
                                QMessageBox.critical(self, "Error", "Incorrect password.")
                                return # Batalkan pembukaan
                            else:
                                # Error pdfium lain
                                QMessageBox.critical(self, "Error", f"Failed to open PDF: {e2}")
                                return
                        except Exception as e_inner:
                            # Error lain saat membuka dengan password
                            QMessageBox.critical(self, "Error", f"Failed to open password-protected file: {e_inner}")
                            return
                    else:
                        # User membatalkan input password
                        return
                else:
                    # Error pdfium lain, bukan masalah password
                    QMessageBox.critical(self, "Error", f"Failed to open PDF: {e}")
                    return
            except Exception as e_general:
                # Error umum lainnya saat load (file korup, dll)
                QMessageBox.critical(self, "Error", f"Failed to open PDF file: {e_general}")
                return

            # Jika berhasil (temp_pdf tidak None)
            self.pdf = temp_pdf
            self.current_file_path = file_path
            self.current_page = 0
            self.setWindowTitle(f"Macan Reader - {os.path.basename(file_path)}")
            self.populate_thumbnails()
            self.go_to_page(self.current_page)
            self.update_ui_state()
            self.add_to_recent_files(file_path)
            
        except Exception as e_outer:
            # Catch-all untuk error tak terduga
            QMessageBox.critical(self, "Error", f"An unexpected error occurred: {e_outer}")
            self.pdf = None
            self.update_ui_state()
    # --- AKHIR PERUBAHAN open_pdf ---

    def close_pdf(self):
        if not self.pdf:
            return
            
        self.pdf.close()
        self.pdf = None
        self.current_file_path = None
        self.current_page = 0
        
        self.setWindowTitle("Macan Reader")
        self.pdf_view_label.clear()
        self.clear_search()
        self.populate_thumbnails()
        self.update_ui_state()
        self.update_status_bar()

    # --- FUNGSI populate_thumbnails DIUBAH UNTUK MENAMBAHKAN NOMOR HALAMAN ---
    def populate_thumbnails(self):
        while self.thumb_layout.count():
            child = self.thumb_layout.takeAt(0)
            if child.widget():
                child.widget().deleteLater()

        if not self.pdf:
            return

        for i in range(len(self.pdf)):
            page = self.pdf[i]
            thumbnail = page.render(scale=0.15).to_pil()
            
            image = QImage(thumbnail.tobytes(), thumbnail.width, thumbnail.height, thumbnail.width * 3, QImage.Format.Format_RGB888)
            pixmap = QPixmap.fromImage(image)
            
            # --- DITAMBAHKAN: Menggambar nomor halaman di atas pixmap ---
            painter = QPainter(pixmap)
            font = painter.font()
            # Buat ukuran font relatif terhadap tinggi thumbnail
            font_size = max(10, int(pixmap.height() / 12))
            font.setPointSize(font_size)
            font.setBold(True)
            painter.setFont(font)
            
            text = str(i + 1)
            metrics = QFontMetrics(font)
            text_rect = metrics.boundingRect(text)
            
            # Posisikan di tengah-bawah
            bg_height = text_rect.height() + 4
            bg_width = text_rect.width() + 8
            bg_x = (pixmap.width() - bg_width) / 2
            bg_y = pixmap.height() - bg_height - 2 # 2px padding dari bawah
            
            bg_rect = QRectF(bg_x, bg_y, bg_width, bg_height)
            
            # Gambar background semi-transparan
            painter.setBrush(QColor(240, 240, 240, 200)) # Putih semi-transparan
            painter.setPen(QColor(100, 100, 100, 150)) # Border abu-abu tipis
            painter.drawRoundedRect(bg_rect, 5, 5)
            
            # Gambar teks nomor halaman
            painter.setPen(QColor(0, 0, 0)) # Teks hitam
            painter.drawText(bg_rect, Qt.AlignmentFlag.AlignCenter, text)
            painter.end()
            # --- AKHIR PENAMBAHAN ---
            
            thumb_label = ThumbnailLabel(i, self.thumb_container)
            thumb_label.setPixmap(pixmap)
            
            thumb_label.page_clicked.connect(self.go_to_page)
            # DITAMBAHKAN: Hubungkan sinyal ekspor
            thumb_label.export_requested.connect(self.export_single_page)
            
            self.thumb_layout.addWidget(thumb_label)
    # --- AKHIR PERUBAHAN populate_thumbnails ---
            
    def display_page(self, page_num):
        if not self.pdf or not (0 <= page_num < len(self.pdf)):
            return

        self.current_page = page_num
        page = self.pdf[page_num]
        
        pil_image = page.render(scale=self.zoom_factor).to_pil()
        
        image = QImage(pil_image.tobytes(), pil_image.width, pil_image.height, pil_image.width * 3, QImage.Format.Format_RGB888)
        pixmap = QPixmap.fromImage(image)
        
        if self.search_results:
            self.draw_search_highlights(page, pixmap)
        
        self.pdf_view_label.setPixmap(pixmap)
        
        self.update_status_bar()
        self.update_ui_state()

    def go_to_page(self, page_num):
        if self.pdf and 0 <= page_num < len(self.pdf):
            self.display_page(page_num)
            
            for i in range(self.thumb_layout.count()):
                widget = self.thumb_layout.itemAt(i).widget()
                if isinstance(widget, ThumbnailLabel):
                    if widget.page_num == page_num:
                        widget.setStyleSheet("border: 2px solid #0078d4; padding: 2px; margin: 5px;")
                    else:
                        widget.setStyleSheet("")

    # --- DITAMBAHKAN: Fungsi untuk dialog Jump to Page ---
    def jump_to_page_dialog(self):
        if not self.pdf:
            return
            
        total_pages = len(self.pdf)
        
        # Minta input integer dari user
        # PERBAIKAN: Hapus keyword 'value=', 'min=', dan 'max='
        # Argument ini harus positional (berurutan)
        page_num, ok = QInputDialog.getInt(self, "Jump to Page",
                                           f"Enter page number (1 - {total_pages}):",
                                           self.current_page + 1, # value (positional)
                                           1,                      # min (positional)
                                           total_pages)            # max (positional)
        
        if ok and page_num:
            # Konversi kembali ke 0-indexed
            self.go_to_page(page_num - 1)

    def set_zoom(self, value):
        self.zoom_factor = value / 100.0
        self.zoom_label.setText(f" {value}% ")
        self.display_page(self.current_page)

    # --- Bagian Logika Pencarian ---

    def clear_search(self):
        self.search_results = []
        self.current_search_index = -1
        if self.pdf:
            self.display_page(self.current_page)
        self.status_bar.clearMessage()

    def run_search(self):
        text = self.search_input.text()
        if not text or not self.pdf:
            self.clear_search()
            return
            
        self.search_results = []
        self.current_search_index = -1
        
        match_case_enabled = self.match_case_check.isChecked()
        
        for i in range(len(self.pdf)):
            page = self.pdf[i]
            # Pastikan hanya memanggil get_textpage() jika memang diperlukan
            try:
                textpage = page.get_textpage()
            except Exception as e:
                # Handle error jika page tidak memiliki textlayer
                print(f"Could not get textpage for page {i}: {e}")
                continue
                
            searcher = textpage.search(text, match_case=match_case_enabled)
            
            # --- PERBAIKAN FINAL: Mengatasi TypeError dan ValueError ---
            while True:
                # 1. Gunakan get_next() untuk mengambil hasil berikutnya (Mengatasi TypeError)
                rect = searcher.get_next() 
                
                if rect is None:
                    # Berhenti ketika tidak ada hasil lagi
                    break
                
                # 2. Cek apakah rect berisi 4 koordinat ([l, b, r, t]) (Mengatasi ValueError)
                # Ini mencegah crash saat mencoba unpack 2 nilai menjadi 4 variabel di line 1038
                if isinstance(rect, (list, tuple)) and len(rect) == 4:
                    self.search_results.append((i, rect))
                else:
                    # Jika formatnya salah (misalnya, hanya 2 nilai), kita abaikan
                    # Ini adalah hasil pencarian yang tidak valid dan penyebab error Anda.
                    print(f"Warning: Skipping non-standard search rect on page {i}: {rect}")
                    pass
            # --- Akhir PERBAIKAN ---
                
        self.status_bar.showMessage(f"Found {len(self.search_results)} result(s).")
        
    def find_next(self):
        text = self.search_input.text()
        if not text: return
        
        if not self.search_results:
            self.run_search()
            
        if not self.search_results:
            return

        self.current_search_index = (self.current_search_index + 1) % len(self.search_results)
        self.highlight_search_result()

    def find_prev(self):
        text = self.search_input.text()
        if not text: return

        if not self.search_results:
            self.run_search()

        if not self.search_results:
            return

        self.current_search_index = (self.current_search_index - 1 + len(self.search_results)) % len(self.search_results)
        self.highlight_search_result()

    def highlight_search_result(self):
        if not self.search_results:
            return
            
        page_num, rect = self.search_results[self.current_search_index]
        
        if page_num != self.current_page:
            self.go_to_page(page_num)
        else:
            self.display_page(page_num)
            
        # --- PERBAIKAN: Gulir (scroll) ke area highlight ---
        page = self.pdf[page_num]
        q_rect = self.convert_pdf_rect_to_pixmap(page, rect)
        
        # Hitung nilai scroll target untuk memusatkan highlight
        h_bar = self.pdf_view_scroll_area.horizontalScrollBar()
        v_bar = self.pdf_view_scroll_area.verticalScrollBar()
        
        target_x = int(q_rect.center().x() - self.pdf_view_scroll_area.viewport().width() / 2)
        target_y = int(q_rect.center().y() - self.pdf_view_scroll_area.viewport().height() / 2)
        
        # Pastikan nilai berada dalam rentang scrollbar
        target_x = max(0, min(target_x, h_bar.maximum()))
        target_y = max(0, min(target_y, v_bar.maximum()))
        
        h_bar.setValue(target_x)
        v_bar.setValue(target_y)
        # --- Akhir PERBAIKAN ---

    def convert_pdf_rect_to_pixmap(self, page, rect):
        l, b, r, t = rect
        page_width, page_height = page.get_size()
        
        pixmap_x = l * self.zoom_factor
        pixmap_y = (page_height - t) * self.zoom_factor
        pixmap_w = (r - l) * self.zoom_factor
        pixmap_h = (t - b) * self.zoom_factor
        
        return QRectF(pixmap_x, pixmap_y, pixmap_w, pixmap_h)

    def draw_search_highlights(self, page, pixmap):
        painter = QPainter(pixmap)
        
        highlight_color = QColor(255, 255, 0, 100)
        active_highlight_color = QColor(255, 165, 0, 150)
        
        painter.setPen(Qt.PenStyle.NoPen)
        
        for i, (p_num, rect) in enumerate(self.search_results):
            if p_num == self.current_page:
                q_rect = self.convert_pdf_rect_to_pixmap(page, rect)
                
                if i == self.current_search_index:
                    painter.setBrush(active_highlight_color)
                else:
                    painter.setBrush(highlight_color)
                    
                painter.drawRect(q_rect)
        
        painter.end()

    # --- Akhir Bagian Logika Pencarian ---

    def save_as_file(self):
        if not self.pdf:
            return
        
        save_path, _ = QFileDialog.getSaveFileName(self, "Save PDF As", "", "PDF Files (*.pdf)")
        if save_path:
            try:
                self.pdf.save(save_path)
                QMessageBox.information(self, "Success", f"File saved successfully to {save_path}")
            except Exception as e:
                QMessageBox.critical(self, "Error", f"Could not save file: {e}")

    # --- Bagian Ekspor ---

    def export_to_png(self):
        if not self.pdf:
            return
        
        base_path, _ = QFileDialog.getSaveFileName(self, "Export Pages as PNG", 
                                                    f"{os.path.splitext(os.path.basename(self.current_file_path))[0]}.png", 
                                                    "PNG Files (*.png)")
        if not base_path:
            return
        
        base_name = os.path.splitext(base_path)[0]
        num_pages = len(self.pdf)
        pad_width = len(str(num_pages))
        
        try:
            for i in range(num_pages):
                page = self.pdf[i]
                # Render dengan kualitas tinggi 300 DPI (skala 300/72)
                img = page.render(scale=300/72.0).to_pil()
                
                file_name = f"{base_name}_page_{str(i+1).zfill(pad_width)}.png"
                img.save(file_name)
                
            QMessageBox.information(self, "Export Successful", 
                                    f"Exported {num_pages} pages to PNG files starting with '{os.path.basename(base_name)}'.")
        except Exception as e:
            QMessageBox.critical(self, "Export Error", f"Failed to export to PNG: {e}")

    def export_to_docx(self):
        if not self.pdf:
            return
        
        save_path, _ = QFileDialog.getSaveFileName(self, "Export as DOCX", 
                                                    f"{os.path.splitext(os.path.basename(self.current_file_path))[0]}.docx", 
                                                    "Word Documents (*.docx)")
        if not save_path:
            return
        
        try:
            import docx
            from docx.shared import Inches
        except ImportError:
            QMessageBox.critical(self, "Missing Dependency",
                                 "Exporting to DOCX requires the 'python-docx' library.\n"
                                 "Please install it using: pip install python-docx")
            return

        try:
            doc = docx.Document()
            
            for i in range(len(self.pdf)):
                page = self.pdf[i]
                # Render pada 150 DPI
                img = page.render(scale=150/72.0).to_pil()
                
                # Simpan gambar ke buffer memori
                img_stream = io.BytesIO()
                img.save(img_stream, format='PNG')
                img_stream.seek(0)
                
                # Tambahkan gambar ke docx, sesuaikan dengan lebar halaman (sekitar 6.5 inci)
                doc.add_picture(img_stream, width=Inches(6.5))
                
                if i < len(self.pdf) - 1:
                    doc.add_page_break()
                        
            doc.save(save_path)
            QMessageBox.information(self, "Export Successful",
                                    f"Successfully exported PDF pages as images to {save_path}.")
        except Exception as e:
            QMessageBox.critical(self, "Export Error", f"Failed to export to DOCX: {e}")

    # --- Slot untuk ekspor halaman tunggal dari context menu ---
    def export_single_page(self, page_num, format_type):
        if not self.pdf or not (0 <= page_num < len(self.pdf)):
            return
        
        base_name = os.path.splitext(os.path.basename(self.current_file_path))[0]
        suggested_name = f"{base_name}_page_{page_num + 1}.{format_type}"
        
        if format_type == "png":
            self.export_page_as_png(page_num, suggested_name)
        elif format_type == "docx":
            self.export_page_as_docx(page_num, suggested_name)
        elif format_type == "pdf":
            self.export_page_as_pdf(page_num, suggested_name)

    def export_page_as_png(self, page_num, suggested_name):
        save_path, _ = QFileDialog.getSaveFileName(self, "Export Page as PNG", 
                                                   suggested_name, "PNG Files (*.png)")
        if not save_path:
            return
        
        try:
            page = self.pdf[page_num]
            img = page.render(scale=300/72.0).to_pil() # 300 DPI
            img.save(save_path)
            QMessageBox.information(self, "Export Successful", f"Page {page_num + 1} saved to {save_path}")
        except Exception as e:
            QMessageBox.critical(self, "Export Error", f"Failed to export page to PNG: {e}")

    def export_page_as_docx(self, page_num, suggested_name):
        save_path, _ = QFileDialog.getSaveFileName(self, "Export Page as DOCX",
                                                   suggested_name, "Word Documents (*.docx)")
        if not save_path:
            return
        
        try:
            import docx
            from docx.shared import Inches
        except ImportError:
            QMessageBox.critical(self, "Missing Dependency",
                                 "Exporting to DOCX requires the 'python-docx' library.\n"
                                 "Please install it using: pip install python-docx")
            return
            
        try:
            doc = docx.Document()
            page = self.pdf[page_num]
            img = page.render(scale=150/72.0).to_pil() # 150 DPI
            
            img_stream = io.BytesIO()
            img.save(img_stream, format='PNG')
            img_stream.seek(0)
            
            doc.add_picture(img_stream, width=Inches(6.5))
            doc.save(save_path)
            QMessageBox.information(self, "Export Successful", f"Page {page_num + 1} saved to {save_path}")
        except Exception as e:
            QMessageBox.critical(self, "Export Error", f"Failed to export page to DOCX: {e}")

    def export_page_as_pdf(self, page_num, suggested_name):
        save_path, _ = QFileDialog.getSaveFileName(self, "Export Page as PDF",
                                                   suggested_name, "PDF Files (*.pdf)")
        if not save_path:
            return
            
        try:
            # Buat PDF kosong baru
            new_pdf = pdfium.PdfDocument.new()
            # Impor halaman tunggal
            new_pdf.import_pages(self.pdf, [page_num])
            # Simpan PDF baru
            new_pdf.save(save_path)
            new_pdf.close()
            QMessageBox.information(self, "Export Successful", f"Page {page_num + 1} saved to {save_path}")
        except Exception as e:
            QMessageBox.critical(self, "Export Error", f"Failed to export page to PDF: {e}")

    # --- Akhir Bagian Ekspor ---

    # --- DIUBAH: Fungsi Print sekarang menggunakan QPrintPreviewDialog ---
    def print_file(self):
        if not self.pdf:
            return

        printer = QPrinter(QPrinter.PrinterMode.HighResolution)
        # Buat dialog print preview
        preview_dialog = QPrintPreviewDialog(printer, self)
        # Hubungkan sinyal paintRequested ke slot yang akan me-render halaman
        preview_dialog.paintRequested.connect(self.handle_paint_request)
        # Tampilkan dialog
        preview_dialog.exec()

    # DITAMBAHKAN: Slot untuk menangani rendering ke printer (untuk preview dan print)
    def handle_paint_request(self, printer):
        painter = QPainter()
        painter.begin(printer)
        
        print_range_setting = printer.printRange()
        from_page = printer.fromPage()
        to_page = printer.toPage()
        
        pages_to_print = list(range(len(self.pdf)))

        if print_range_setting == QPrinter.PrintRange.PageRange:
            pages_to_print = list(range(from_page - 1, min(to_page, len(self.pdf))))
        elif print_range_setting == QPrinter.PrintRange.CurrentPage:
            pages_to_print = [self.current_page]

        for i, page_idx in enumerate(pages_to_print):
            if page_idx >= len(self.pdf):
                continue
                
            page = self.pdf[page_idx]
            
            dpi = printer.resolution()
            if dpi == 0:
                dpi = 300
            scale = dpi / 72.0
            
            img = page.render(scale=scale).to_pil()
            
            q_img = QImage(img.tobytes(), img.width, img.height, img.width * 3, QImage.Format.Format_RGB888)
            
            page_rect = printer.pageRect(QPrinter.Unit.DevicePixel)
            img_pixmap = QPixmap.fromImage(q_img)
            
            # --- PERBAIKAN: Konversi QSizeF ke QSize menggunakan .toSize() ---
            scaled_pixmap = img_pixmap.scaled(page_rect.size().toSize(), Qt.AspectRatioMode.KeepAspectRatio, Qt.TransformationMode.SmoothTransformation)
            # ----------------------------------------------------------------
            
            x = (page_rect.width() - scaled_pixmap.width()) / 2
            y = (page_rect.height() - scaled_pixmap.height()) / 2
            
            painter.drawPixmap(int(x), int(y), scaled_pixmap)
            
            if i < len(pages_to_print) - 1:
                printer.newPage()
        painter.end()

    # --- CHANGED: Help content translated to English ---
    def show_help_content(self):
        help_text = """
        <h2>Macan Reader - Help</h2>
        
        <h3>Navigation</h3>
        <ul>
            <li><b>Next Page:</b> Click the 'Next' arrow button or press the <b>Right Arrow</b> key.</li>
            <li><b>Previous Page:</b> Click the 'Previous' arrow button or press the <b>Left Arrow</b> key.</li>
            <li><b>Go to Page:</b> Click a thumbnail in the left-hand pane.</li>
            <li><b>Jump to Page (Ctrl+G):</b> Use the 'Jump' (<b>#</b>) button or <b>Edit > Jump to Page</b> to enter a specific page number.</li>
        </ul>

        <h3>Zooming & Panning</h3>
        <ul>
            <li><b>Zoom In/Out:</b> Use the <b>+</b> and <b>-</b> buttons, the slider in the status bar, or <b>Ctrl+</b>/<b>Ctrl-</b>.</li>
            <li><b>Reset Zoom:</b> Click the 'target' icon or press <b>Ctrl+0</b> to reset the zoom to 100%.</li>
            <li><b>Pan (Move):</b> Click the 'Hand' icon to enter Pan mode. Click and drag the page to move the view. Click the icon again to exit.</li>
        </ul>

        <h3>File</h3>
        <ul>
            <li><b>Open (Ctrl+O):</b> Open a new PDF file. If the file is password-protected, you will be prompted to enter it.</li>
            <li><b>Save As:</b> Save a copy of the currently open PDF.</li>
            <li><b>Print (Ctrl+P):</b> Open the print preview to print the document.</li>
            <li><b>Export As:</b> Export all pages of the document as PNG images or as images within a DOCX file.</li>
            <li><b>Recent Files:</b> Open a recently viewed file.</li>
        </ul>
        
        <h3>Search (Ctrl+F)</h3>
        <ul>
            <li>Type your keyword in the 'Find' bar at the top.</li>
            <li>Click <b>Next</b> or <b>Previous</b> to navigate through the results.</li>
            <li>Check <b>Match Case</b> for case-sensitive searching.</li>
            <li>The current result is highlighted in orange; other results are highlighted in yellow.</li>
        </ul>

        <h3>Thumbnails</h3>
        <ul>
            <li><b>Show/Hide:</b> Use <b>Window > Show Thumbnail Pane</b> or the 'X' button on the pane itself.</li>
            <li><b>Page Numbers:</b> Each thumbnail displays its page number at the bottom.</li>
            <li><b>Context Menu:</b> Right-click on a thumbnail to export that specific page as a PDF, PNG, or DOCX.</li>
        </ul>

        <h3>Themes</h3>
        <ul>
            <li>Change the application's appearance from the <b>Themes</b> menu. Your choice will be saved.</li>
        </ul>
        """
        
        help_dialog = QMessageBox(self)
        help_dialog.setWindowTitle("Macan Reader Help")
        help_dialog.setIcon(QMessageBox.Icon.Question)
        
        help_dialog.setText(" ") 
        
        scroll_area = QScrollArea(help_dialog)
        scroll_area.setWidgetResizable(True)
        scroll_area.setMinimumSize(500, 300)
        
        help_label = QLabel(help_text, scroll_area)
        help_label.setTextFormat(Qt.TextFormat.RichText)
        help_label.setWordWrap(True)
        help_label.setAlignment(Qt.AlignmentFlag.AlignTop)
        
        scroll_area.setWidget(help_label)
        
        layout = help_dialog.layout()
        
        old_widget = layout.itemAtPosition(0, 1).widget()
        if old_widget:
            old_widget.deleteLater()
            
        layout.addWidget(scroll_area, 0, 1, 1, 1)
        
        help_dialog.exec()

    def show_about_dialog(self):
        about_text = """
        <h2>Macan Reader</h2>
        <p>A professional PDF reader built with PySide6 and PyPDFium2.</p>
        <p>Fast, efficient, and reliable.</p>
        <p>v3.4.0</p>
        <br>
        <p>© 2025 - Danx Exodus - Macan Angkasa</p>
        """
        QMessageBox.about(self, "About Macan Reader", about_text)

    def add_to_recent_files(self, file_path):
        recent_files = self.settings.value("recentFiles", [])
        if file_path in recent_files:
            recent_files.remove(file_path)
        recent_files.insert(0, file_path)
        del recent_files[10:]
        self.settings.setValue("recentFiles", recent_files)
        self.update_recent_files_menu()
        
    def update_recent_files_menu(self):
        self.recent_files_menu.clear()
        recent_files = self.settings.value("recentFiles", [])
        
        if recent_files:
            for file_path in recent_files:
                action = QAction(os.path.basename(file_path), self)
                action.setData(file_path)
                action.triggered.connect(self.open_recent_file)
                self.recent_files_menu.addAction(action)
            self.recent_files_menu.addSeparator()
            self.recent_files_menu.addAction(self.clear_recent_action)
        else:
            no_recent_action = QAction("No Recent Files", self)
            no_recent_action.setEnabled(False)
            self.recent_files_menu.addAction(no_recent_action)

    def clear_recent_files(self):
        self.settings.setValue("recentFiles", [])
        self.update_recent_files_menu()

    def open_recent_file(self):
        action = self.sender()
        if action:
            file_path = action.data()
            if os.path.exists(file_path):
                self.open_pdf(file_path)
            else:
                QMessageBox.warning(self, "File Not Found", f"The file '{file_path}' could not be found.")
                recent_files = self.settings.value("recentFiles", [])
                if file_path in recent_files:
                    recent_files.remove(file_path)
                    self.settings.setValue("recentFiles", recent_files)
                    self.update_recent_files_menu()

    def update_ui_state(self):
        is_pdf_loaded = self.pdf is not None
        
        self.save_as_action.setEnabled(is_pdf_loaded)
        self.print_action.setEnabled(is_pdf_loaded)
        self.close_action.setEnabled(is_pdf_loaded)
        # DITAMBAHKAN: Enable/Disable aksi ekspor
        self.export_png_action.setEnabled(is_pdf_loaded)
        self.export_docx_action.setEnabled(is_pdf_loaded)
        
        self.find_action.setEnabled(is_pdf_loaded)
        self.jump_to_page_action.setEnabled(is_pdf_loaded) # DITAMBAHKAN
        self.pan_action.setEnabled(is_pdf_loaded)
        if not is_pdf_loaded and self.pan_action.isChecked():
            self.pan_action.setChecked(False)
            self.toggle_pan_mode(False)
        
        self.zoom_in_action.setEnabled(is_pdf_loaded)
        self.zoom_out_action.setEnabled(is_pdf_loaded)
        self.zoom_reset_action.setEnabled(is_pdf_loaded)
        self.zoom_slider.setEnabled(is_pdf_loaded)
        
        if is_pdf_loaded:
            self.prev_page_action.setEnabled(self.current_page > 0)
            self.next_page_action.setEnabled(self.current_page < len(self.pdf) - 1)
        else:
            self.prev_page_action.setEnabled(False)
            self.next_page_action.setEnabled(False)

    def update_status_bar(self):
        if self.pdf:
            self.file_info_label.setText(os.path.basename(self.current_file_path))
            self.page_info_label.setText(f" Page: {self.current_page + 1} of {len(self.pdf)} ")
        else:
            self.file_info_label.setText("No file opened")
            self.page_info_label.setText("")

    def dragEnterEvent(self, event):
        if event.mimeData().hasUrls():
            url = event.mimeData().urls()[0]
            if url.isLocalFile() and url.toLocalFile().lower().endswith('.pdf'):
                event.acceptProposedAction()

    def dropEvent(self, event):
        url = event.mimeData().urls()[0]
        file_path = url.toLocalFile()
        self.open_pdf(file_path)

    def closeEvent(self, event):
        if self.pdf:
            self.close_pdf()
        self.settings.sync()
        event.accept()

if __name__ == "__main__":
    app = QApplication(sys.argv)
    
    file_to_open = None
    if len(sys.argv) > 1:
        path = sys.argv[1]
        if os.path.exists(path) and path.lower().endswith('.pdf'):
            file_to_open = path
            
    window = MacanReader(file_to_open)
    window.show()
    sys.exit(app.exec())